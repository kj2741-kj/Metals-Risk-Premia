"""
generate_research_appendix.py
Generates a SEPARATE research-appendix Word document for LME Copper Risk Premia.
It documents every strategy variant and method we explored but did NOT make the
production default / surface prominently in the live dashboard - the "search"
behind the dashboard's final choices.

Numbers are active-day annualised Sharpe on the full sample (2006-01-03 -> 2025-12-31,
~5,053 trading days), computed with the dashboard's exact signal+metric logic.
Output: LME_Copper_RiskPremia_ResearchAppendix.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)

COPPER = RGBColor(0xB8, 0x73, 0x33); DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY   = RGBColor(0x55, 0x55, 0x55); LIGHT = RGBColor(0x78, 0x78, 0x78)

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = COPPER
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4"); bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "B87333")
    pBdr.append(bot); pPr.append(pBdr); return p

def heading2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK; return p

def body(text, space_after=4, italic=False, colour=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.font.size = Pt(9.5); r.italic = italic
    if colour: r.font.color.rgb = colour
    return p

def mixed_bullet(parts):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2); p.paragraph_format.left_indent = Inches(0.2)
    for text, bold in parts:
        r = p.add_run(text); r.font.size = Pt(9.5); r.bold = bold
    return p

def space(pts=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts)

def add_table(headers, rows, widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (h, w) in enumerate(zip(headers, widths)):
        c = t.rows[0].cells[i]; c.width = Inches(w); set_cell_bg(c, "B87333")
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        fill = "F5EFE8" if ri % 2 == 0 else "FAFAFA"
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]; c.width = Inches(widths[ci]); set_cell_bg(c, fill)
            r = c.paragraphs[0].add_run(val); r.font.size = Pt(9)
            if ci == 0: r.bold = True
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
    space(6); return t

# ── Cover ─────────────────────────────────────────────────────────────────────
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("LME Copper Risk Premia"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = COPPER
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Research Appendix - Strategies & Methods Explored (Not Production Defaults)")
r.font.size = Pt(12); r.font.color.rgb = DARK
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Kartavya Joshi  |  Companion to the Project Overview & live dashboard"); r.font.size = Pt(9.5); r.font.color.rgb = LIGHT
space(8)

# ── 1. Purpose ─────────────────────────────────────────────────────────────────
heading1("1. Purpose & Scope")
body("The live dashboard presents only the surviving, default configuration of each signal family. "
     "This appendix records the full exploration behind those choices: every variant tested, every "
     "method considered, and the sensitivity/robustness studies run - including dead-ends and "
     "configurations that were demoted or kept research-only. It is the 'search' behind the dashboard.")
body("Convention: all Sharpe figures are active-day annualised (mean/std x sqrt(252) over days with a "
     "non-zero position), full sample 2006-2025, computed with the dashboard's exact signal and metric "
     "logic. Signal from F1_raw (or curve contracts); P&L always on F1_continuous. 'Same-Day' = shift 1 "
     "(trade at signal close, no look-ahead); 'Lag-1' = shift 2 (one further day of execution delay).",
     italic=True, colour=GREY)

# ── 2. Momentum ────────────────────────────────────────────────────────────────
heading1("2. Momentum - Variants Explored")
body("A broad MA-crossover pair scan (5 <= m < n <= 200) and the Baz-Granger CTA family were tested. "
     "MA(35,43) was the max-IS-Sharpe pair and is the portfolio momentum leg; the CTA variants are "
     "exposed in the dashboard dropdown but are not the default.")
add_table(
    ["Variant", "Same-Day", "Lag-1", "Status"],
    [["MA(35,43) crossover", "+0.72", "+0.63", "DEFAULT (portfolio leg)"],
     ["MA pair scan 5<=m<n<=200", "-", "-", "MA(35,43) = scan winner"],
     ["CTA Paper 3-timescale (8/16/32 vs 24/48/96)", "+0.11", "+0.08", "Faithful to Baz-Granger; weak for copper - reference only"],
     ["CTA single (9,21)", "+0.45", "+0.33", "Explored (dropdown)"],
     ["CTA single (8,21)", "+0.43", "-", "Explored (dropdown)"],
     ["CTA single (10,19)", "-", "+0.34", "Explored (dropdown)"],
     ["CTA single (14,15)", "-", "+0.35", "Explored (dropdown)"],
     ["Anchors EW + IS-opt weights (walk-forward)", "-", "-", "Explored; optimiser concentrates on MA(35,43)"]],
    [2.6, 0.9, 0.8, 2.1])
body("Takeaway: the literature-faithful CTA-paper signal underperforms a simple MA(35,43) crossover for "
     "copper (0.11 vs 0.72). The medium-frequency MA crossover is the production momentum signal; the "
     "CTA machinery is retained for completeness and auditing of the paper methodology.", italic=True)

# ── 3. Carry ───────────────────────────────────────────────────────────────────
heading1("3. Carry - Variants Explored")
body("Five carry families were tested. The 20-day change in the (F1-F2)/F1 roll yield (curve momentum) "
     "is the production leg; level, z-score, long-slope and alternative horizons were explored and demoted.")
add_table(
    ["Signal", "Sharpe (Same-Day)", "Status"],
    [["Carry-Momentum 20d  (delta-20d of (F1-F2)/F1)", "+0.52", "DEFAULT (portfolio leg)"],
     ["Carry-Momentum 60d", "+0.39", "Too slow - demoted"],
     ["Carry-Momentum 1-day", "(high, fragile)", "Microstructure artifact - collapses at higher TC and under shift-2; REJECTED"],
     ["(F1-F2)/F1 roll-yield level", "+0.10", "Honest baseline - weak"],
     ["(F1-F3)/F1 / (Cash-3M)/Cash levels", "~+0.1", "Explored - similar to F1-F2 level"],
     ["Z-score 252d of roll yield", "+0.26", "Explored (dropdown)"],
     ["Long-slope (Fj-Fk)/Fk, F3-F15..F12-F24", "+0.14 (best pair)", "Explored - weak"],
     ["Legacy shift-0 'same-day' level carry", "~+0.62", "LOOK-AHEAD - removed (booked an already-realised move)"]],
    [3.0, 1.3, 2.1])
body("Takeaway: the headline ~0.62 'same-day' carry was a look-ahead artifact (shift 0). The honest level "
     "carry is only ~0.10; the 20-day roll-yield momentum (+0.52) is the genuine, TC- and lag-robust "
     "survivor and is the portfolio carry leg.", italic=True)

# ── 4. Value ───────────────────────────────────────────────────────────────────
heading1("4. Value - Variants Explored")
body("A full contract scan (F1-F15) x lookback grid was run for V1 MA-reversion, plus the V2 Baz-Granger "
     "reversal across lookbacks. F8 at a 5yr lookback is the copper-optimal V1 contract; the NGL energy "
     "paper's F12 reference is sub-optimal for copper. V2 is event-driven and was demoted.")
add_table(
    ["Variant", "Sharpe", "Status"],
    [["V1 F8, 5yr, +/-10%", "+0.28 SD / +0.33 L1", "DEFAULT (portfolio leg); robust OOS (+0.43)"],
     ["V1 F12, 5yr (NGL energy reference)", "+0.235 L1", "NGL paper tenor - sub-optimal for copper"],
     ["V1 contract scan F1-F15", "-", "F8 = copper-optimal contract"],
     ["V1 F14, 7yr / 10yr", "-0.06 / +0.03", "Weak - demoted"],
     ["V2 BG reversal, 3yr", "+0.26", "Explored"],
     ["V2 BG reversal, 10yr", "+0.37 (best IS)", "Highest IS but fragile - OOS collapses to +0.07 (COVID-concentrated); optional only"],
     ["V2 BG reversal, 1/5/7yr", "neg / trap / ~0.16", "Non-monotonic by lookback; 5yr is a trap"],
     ["Threshold +/-10% vs +/-15-20%", "-", "+/-10% from NGL energy calibration; wider band suggested for lower-vol copper - untested toggle"]],
    [2.9, 1.5, 2.0])
body("Takeaway: copper's value edge lives at F8/5yr, not the NGL energy paper's F12. V2 BG 10yr posts the "
     "highest in-sample Sharpe but its edge is concentrated in the 2020-2021 dislocation and does not "
     "persist out-of-sample, so V1 F8 (robust OOS) is the production value sleeve.", italic=True)

# ── 5. Portfolio weighting ──────────────────────────────────────────────────────
heading1("5. Portfolio Weighting Research")
body("Equal-weight (1/3 each) is the production default. Inverse-volatility weighting was studied as an "
     "alternative; we swept the trailing return-vol estimation window. A 63-day window is Sharpe-optimal, "
     "but equal-weight still dominates on net Sharpe and drawdown - so the trailing-window sweep is "
     "research-only and not surfaced on the dashboard (which keeps a single fixed 63d inverse-vol toggle).")
add_table(
    ["Weighting", "Gross Sh", "Net 5bps", "Net 10bps", "Ann %", "Max DD ($/MT)"],
    [["Equal-Weight (1/3 each)", "1.045", "0.959", "0.899", "11.9", "-2,020"],
     ["Inverse-Vol - 10d", "0.858", "0.800", "0.743", "13.3", "-3,064"],
     ["Inverse-Vol - 21d", "0.908", "0.852", "0.797", "13.7", "-3,040"],
     ["Inverse-Vol - 42d", "0.903", "0.848", "0.794", "13.3", "-3,054"],
     ["Inverse-Vol - 63d (selected)", "0.941", "0.887", "0.833", "13.5", "-3,031"],
     ["Inverse-Vol - 126d", "0.906", "0.852", "0.800", "12.4", "-3,038"],
     ["Inverse-Vol - 189d", "0.910", "0.858", "0.807", "12.2", "-3,031"],
     ["Inverse-Vol - 252d", "0.884", "0.832", "0.782", "11.7", "-3,029"]],
    [2.0, 0.85, 0.85, 0.9, 0.7, 1.1])
mixed_bullet([("Regime-conditional weighting: ", True), ("proposed - downweight carry in contango regimes, "
              "upweight value in dislocations (high VIX + low positioning z-score). NOT yet built.", False)])
mixed_bullet([("Walk-forward design: ", True), ("IS = 5yr (1,260d) rolling, OOS = 1yr (252d), "
              "non-overlapping; used for all OOS Sharpe estimates.", False)])

# ── 6. Methodology & robustness ─────────────────────────────────────────────────
heading1("6. Methodology & Robustness Studies")
heading2("6.1  Execution timing & look-ahead removal")
body("Three conventions were compared. Shift 0 ('legacy same-day') books the t-1->t move that had already "
     "happened when the signal became known - pure look-ahead - and was removed. It had inflated the level "
     "carry Sharpe from ~0.10 (honest) to ~0.62. Production uses shift 1 (Same-Day, trade at signal close, "
     "first return t->t+1) or shift 2 (Lag-1, one further day of delay).")
heading2("6.2  Warmup handling sensitivity")
body("Rolling/EWMA warmup must require full windows (warmup days flat, excluded from active-day Sharpe). "
     "A naive min_periods=1 warmup lets the signal trade during initialisation on under-estimated vol; this "
     "INFLATES the CTA strategies by roughly +0.05 to +0.10 Sharpe (their warmup is ~314d: 63 for sigma-63 "
     "then 252 for sigma-252 of y). The dashboard uses the correct full-window logic; the inflated figures "
     "are what a mishandled warmup would have shown. (MA strategies are warmup-invariant - 43d warmup.)")
heading2("6.3  Transaction-cost treatment")
body("TC_Cost(t) = |Position(t) - Position(t-1)| x (tc_bps / 10,000 / 2) x F1_Continuous(t). The /2 is the "
     "half-spread per side (tc_bps is round-trip). F1_continuous (not F1_raw) is used for unit-consistency: "
     "the entire P&L, return and drawdown stack is denominated in F1_continuous, so the cost subtracted from "
     "it must be on the same price scale (the two series diverge historically under ratio back-adjustment). "
     "Base case 5bps; sensitivity run at 5 / 10 / 20 bps - slow signals (carry, value) are cheap, while "
     "higher-turnover and daily-reweighted (inverse-vol) variants degrade faster with TC.")

# ── 7. Decision log ─────────────────────────────────────────────────────────────
heading1("7. Decision Log - Production vs Research-Only")
add_table(
    ["Item", "In live dashboard?", "Note"],
    [["MA(35,43) momentum", "Yes (default)", "Portfolio momentum leg"],
     ["CTA paper / CTA single variants", "Yes (dropdown)", "Not default - weaker for copper"],
     ["Carry-Momentum 20d", "Yes (default)", "Portfolio carry leg"],
     ["Carry level / z-score / long-slope", "Yes (dropdown)", "Not default - demoted"],
     ["Carry-Mom 1d / 60d", "No", "Rejected (artifact / too slow)"],
     ["Value V1 F8 5yr", "Yes (default)", "Portfolio value leg"],
     ["Value V1 F12 (NGL ref) & F1-F15 scan", "Yes (dropdown/chart)", "Context - F8 wins for copper"],
     ["Value V2 BG reversal", "Yes (optional)", "Demoted - fragile OOS"],
     ["EW portfolio", "Yes (default)", "Best net Sharpe + drawdown"],
     ["Inverse-vol 63d", "Yes (toggle)", "Single fixed window only"],
     ["Inverse-vol trailing-window sweep", "No", "Research-only (this doc + Project Overview 6.1)"],
     ["Regime-conditional weighting", "No", "Proposed, not built"],
     ["Warmup-sensitivity / look-ahead studies", "No", "Methodology - this doc"]],
    [2.6, 1.6, 2.4])

body("This appendix is a living research log; numbers correspond to the run dated in the file header and "
     "should be re-derived from the standalone scripts when the data is refreshed.", italic=True, colour=LIGHT)

out = r"C:\Users\Kartavya\Metals Risk Premia\LME_Copper_RiskPremia_ResearchAppendix.docx"
doc.save(out)
print(f"Saved: {out}")
