"""
generate_findings_synthesis_doc.py
Synthesis of our LME Copper risk-premia findings, the live dashboard, and -- per
Prof. Ilia's request -- a focus on more recent behaviour. Also links our work to
Gorton-Hayashi-Rouwenhorst (2013) and states our view on fundamentals.
Output: LME_Copper_RiskPremia_Findings_Synthesis.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_REPO_ROOT = os.path.dirname(_SCRIPT_DIR)
DOCS_DIR = os.path.join(_REPO_ROOT, "docs")

doc=Document()
s=doc.sections[0]; s.top_margin=Cm(2.0); s.bottom_margin=Cm(2.0); s.left_margin=Cm(2.2); s.right_margin=Cm(2.2)
COPPER=RGBColor(0xB8,0x73,0x33); DARK=RGBColor(0x1A,0x1A,0x1A); GREY=RGBColor(0x55,0x55,0x55); LIGHT=RGBColor(0x78,0x78,0x78)
def cell_bg(c,h):
    p=c._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:color"),"auto"); sh.set(qn("w:fill"),h); p.append(sh)
def h1(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t.upper()); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=COPPER
    pr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
    bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"4"); bt.set(qn("w:space"),"1"); bt.set(qn("w:color"),"B87333"); b.append(bt); pr.append(b)
def h2(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(10.5); r.font.color.rgb=DARK
def body(t,sa=5,it=False,col=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(sa); r=p.add_run(t); r.font.size=Pt(10.5); r.italic=it
    if col: r.font.color.rgb=col
def bullet(t):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Inches(0.25)
    r=p.add_run(t); r.font.size=Pt(10.5)
def kv(label,text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Inches(0.25)
    r=p.add_run(label+": "); r.bold=True; r.font.size=Pt(10.5); r2=p.add_run(text); r2.font.size=Pt(10.5)
def space(p=4): doc.add_paragraph().paragraph_format.space_after=Pt(p)
def table(headers,rows,widths):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(hh,w) in enumerate(zip(headers,widths)):
        c=t.rows[0].cells[i]; c.width=Inches(w); cell_bg(c,"B87333")
        r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    for ri,row in enumerate(rows):
        fill="F5EFE8" if ri%2==0 else "FAFAFA"
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.width=Inches(widths[ci]); cell_bg(c,fill)
            r=c.paragraphs[0].add_run(v); r.font.size=Pt(9)
            if ci==0: r.bold=True
            c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER
    space(6)

# Cover
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("LME Copper Risk Premia"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=COPPER
st=doc.add_paragraph(); st.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=st.add_run("Findings Synthesis, the Live Dashboard, and Recent Behaviour"); r.font.size=Pt(12); r.font.color.rgb=DARK
a=doc.add_paragraph(); a.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=a.add_run("Kartavya Joshi  |  Prepared in response to Prof. Ilia Bouchoev's note"); r.font.size=Pt(9.5); r.font.color.rgb=LIGHT
space(6)

h1("1. Purpose")
body("This note synthesises what we have learned from the LME copper risk-premia analysis, what the live "
     "dashboard currently shows, and - as Prof. Ilia asked - focuses on more recent behaviour. It also "
     "connects our results to the inventory paper he sent (Gorton-Hayashi-Rouwenhorst 2013, summarised "
     "separately) and states our view on fundamentals. Performance figures are active-day annualised "
     "Sharpe on LME copper, 2006-2025, computed with the dashboard's exact logic (signals from F1_raw, "
     "P&L on the roll-adjusted F1_continuous, transaction costs charged on the actual traded F1_raw price).")

h1("2. What We Built")
kv("Three single-factor sleeves","Momentum (MA(35,43) crossover on front price), Carry (20-day change in "
   "the (F1-F2)/F1 roll yield), and Value (price vs its 5-year mean on the F8 contract, +/-10% band).")
kv("A combined portfolio","Equal-weight 1/3 each (default), with an inverse-volatility alternative.")
kv("Honest accounting","No look-ahead (signal at close t trades t->t+1); transaction costs on F1_raw; "
   "rolling windows fully initialised before trading; walk-forward out-of-sample testing (5yr IS / 1yr OOS).")
kv("Deliverables","A 10-tab Streamlit dashboard (live), full Excel tradebooks per signal and the portfolio, "
   "and research docs (this note, a methods appendix, and the GHR paper summary).")

h1("3. Headline Results (Full Sample, 2006-2025)")
table(["Sleeve / Portfolio","Gross Sharpe","Net 5bps","Comment"],
    [["Momentum MA(35,43)","0.72","0.70","Steadiest standalone sleeve"],
     ["Carry CarryMom 20d","0.52","0.46","Curve-momentum; best of the carry family"],
     ["Value V1 F8 5yr","0.28","0.27","Weakest, regime-dependent"],
     ["Portfolio - Equal-Weight","1.05","0.96","Diversification lifts it above every sleeve"],
     ["Portfolio - Inverse-Vol 63d","0.94","0.89","Higher return, worse drawdown; EW preferred"]],
    [2.6,1.1,1.0,2.6])
body("The central result is the diversification gain: three weakly-correlated sleeves combine into a "
     "portfolio Sharpe (~1.0 net) well above any single sleeve - the classic carry+momentum+value benefit.")

h1("4. Recent Behaviour  (Prof. Ilia's Focus)")
body("Splitting the history into sub-periods is where the interesting story lives. The sleeves rotate "
     "leadership markedly, and recent behaviour differs sharply from the long-run averages.")
table(["Sleeve / Portfolio (Gross Sharpe)","Pre-2020","2020-21","2022","2023-25","Last 252d"],
    [["Momentum MA(35,43)","0.69","1.17","0.30","0.75","0.89"],
     ["Carry CarryMom 20d","0.55","0.21","0.92","0.47","1.95"],
     ["Value V1 F8 5yr","0.46","-1.05","2.08","0.15","-1.19"],
     ["Portfolio Equal-Weight","1.11","0.42","1.52","1.15","2.17"]],
    [3.0,1.2,1.1,1.0,1.0,1.1][:6])
h2("Key recent observations")
kv("Carry is the standout right now","Carry-momentum's last-12-month Sharpe (~1.95) is its best of the whole "
   "sample - the copper curve has given clear, tradable backwardation/contango swings lately.")
kv("Momentum is quietly solid","~0.89 over the last year and strongest during the 2020-21 trend; a reliable "
   "all-weather contributor.")
kv("Value is detracting","After a superb 2022 (+2.08, fading the post-COVID spike), value has been negative "
   "recently (~ -1.2 last year) - copper has stayed 'expensive' versus its 5-year mean and kept rising, so "
   "the mean-reversion bet has bled. This is the regime-conditional weakness we have flagged throughout.")
kv("Yet the portfolio is at its best","Despite value's drag, the equal-weight portfolio posted its strongest "
   "sub-period (~2.1 last year) - carry and momentum more than offset value. This is the diversification "
   "thesis working in real time: the sleeves are doing different things at different times.")
kv("Current positioning is net-SHORT copper","At end-2025 the book is split: Momentum +1 (uptrend), Carry -1 "
   "(contango), Value -1 (expensive vs mean) -> equal-weight net position about -1/3. Over the last year the "
   "average absolute portfolio position was only ~0.37, i.e. the sleeves frequently disagree and partially "
   "net out - a feature, not a bug, of combining diversifying signals.")

h1("5. Per-Sleeve Synthesis")
h2("5.1  Momentum")
bullet("Medium-frequency MA(35,43) beats the literature CTA signal on copper (0.72 vs ~0.11) - the simple "
       "crossover is the workhorse.")
bullet("Most valuable in trending regimes (2020-21); weakest in choppy 2022. No flat zone - always in the market.")
h2("5.2  Carry")
bullet("The 20-day change in roll yield (curve momentum) clearly beats the raw carry level (~0.52 vs ~0.10). "
       "Trading the CHANGE in the curve, not its level, is what works.")
bullet("Recently the strongest sleeve; historically the best risk-adjusted single signal after momentum.")
h2("5.3  Value")
bullet("Copper-optimal contract is F8/5yr (not the F12 from the energy literature). Weak and regime-dependent: "
       "great when a dislocation reverts (2022), painful when copper trends rich (recent).")
bullet("Kept as a diversifier, not a return driver. We deliberately use the robust F8 version over the "
       "higher-in-sample-but-fragile Baz-Granger 10yr reversal.")

h1("6. Methodology Notes (so the numbers are trustworthy)")
kv("No look-ahead","A removed legacy 'same-day' convention had inflated carry to ~0.62; the honest level "
   "carry is ~0.10. All figures here use the corrected timing.")
kv("Transaction costs on F1_raw","Costs are charged on the actual traded front-month price, not the "
   "back-adjusted index; for copper this is a tiny, positive adjustment (curves roll ~1%).")
kv("Warmup handled","Rolling/EWMA windows must be fully initialised before trading; metrics use active days only.")
kv("Out-of-sample","Walk-forward (5yr in-sample / 1yr out-of-sample, non-overlapping) underpins the OOS Sharpes.")

h1("7. Link to the Inventory Paper (GHR 2013)")
body("Gorton-Hayashi-Rouwenhorst show that commodity risk premiums are driven by physical inventories, and "
     "- key for us - that price-based signals (the basis/carry, momentum, volatility) are reliable proxies "
     "for the inventory state. Their cross-sectional sorts on basis and on momentum earn ~10-12%/yr; the two "
     "are 0.85 correlated because both select scarce, low-inventory commodities.")
kv("Why this matters","Our carry and momentum sleeves are exactly those price proxies. The paper provides the "
   "economic licence for why they work - they harvest the inventory-driven risk premium without needing "
   "inventory data.")
kv("On positioning","GHR find trader positions (CFTC hedging pressure) do NOT predict returns - so we should "
   "treat any COT-style data as context, not a forecasting edge.")

h1("8. On Fundamentals - We Agree with Prof. Ilia")
body("Prof. Ilia's low expectations for metals fundamentals on public data are well supported by GHR's own "
     "evidence: the inventory->basis and inventory->return links are statistically strong for hard-to-store "
     "groups (energies, grains) but WEAKEST for easy-to-store METALS. Metals have large 'normal' inventories "
     "and rarely stock out, so the convenience-yield curve is flat except in genuine shortages.")
kv("Our stance","Keep fundamentals exploratory and low-expectation for copper. If inventory ever adds value, "
   "GHR implies it will only be in the convex, steep regime - stocks far below normal - i.e. an event/regime "
   "overlay, not a smooth daily factor. Worth a brief discussion, not a major build.")

h1("9. Open Questions & Next Steps")
bullet("Prof. Ilia mentioned a related idea to discuss - happy to scope it against the current framework.")
bullet("Recent leadership rotation (carry strong, value weak) suggests revisiting whether a light regime "
       "or trend filter on the value sleeve would help - to be weighed against over-fitting risk.")
bullet("Possible extension of the same price-signal framework to the other LME base metals (the original "
       "10-metal goal), where the cross-section may add a genuine relative-value dimension.")
bullet("A short, honest note on inventory as an event-overlay for copper, drawing on GHR, if of interest.")

body("Companion documents: GHR_2013_Inventory_Paper_Summary.docx (plain-language paper summary); "
     "LME_Copper_RiskPremia_ResearchAppendix.docx (methods & strategies explored).", sa=2, col=LIGHT)

out=os.path.join(DOCS_DIR, "LME_Copper_RiskPremia_Findings_Synthesis.docx")
doc.save(out); print("Saved:",out)
