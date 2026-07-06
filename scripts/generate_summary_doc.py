"""
Generates the Baz-Granger paper summary as a Word document.
Run: python generate_summary_doc.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_REPO_ROOT = os.path.dirname(_SCRIPT_DIR)
DOCS_DIR = os.path.join(_REPO_ROOT, "docs")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
COPPER     = RGBColor(0xB8, 0x73, 0x33)   # copper accent
CHARCOAL   = RGBColor(0x1E, 0x1E, 0x1E)   # near-black body
DARK_GREY  = RGBColor(0x44, 0x44, 0x44)   # subheading
MID_GREY   = RGBColor(0x88, 0x88, 0x88)   # caption / note
TABLE_HEAD = RGBColor(0x2C, 0x2C, 0x2C)   # table header bg (set via xml)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: set paragraph font defaults ──────────────────────────────────────
def _set_run(run, size=11, bold=False, italic=False, colour=CHARCOAL,
             font_name="Calibri"):
    run.font.name  = font_name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = colour

def body(para, text, size=11, bold=False, italic=False, colour=CHARCOAL,
         font_name="Calibri"):
    run = para.add_run(text)
    _set_run(run, size=size, bold=bold, italic=italic, colour=colour,
             font_name=font_name)
    return run

# ── Cover page ────────────────────────────────────────────────────────────────
doc.add_paragraph()   # breathing room
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
body(cover_title,
     "Dissecting Investment Strategies\nin the Cross Section and Time Series",
     size=24, bold=True, colour=COPPER, font_name="Calibri Light")

doc.add_paragraph()
cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
body(cover_sub,
     "A Practitioner Summary with Mathematics, Examples, and Application to Metals",
     size=13, italic=True, colour=DARK_GREY)

doc.add_paragraph()
cover_auth = doc.add_paragraph()
cover_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
body(cover_auth,
     "Source: Baz, Granger, Harvey, Le Roux, Rattray — Man AHL (December 2015)",
     size=11, colour=MID_GREY)

cover_auth2 = doc.add_paragraph()
cover_auth2.alignment = WD_ALIGN_PARAGRAPH.CENTER
body(cover_auth2, "Summary by: Kartavya Joshi", size=11, colour=MID_GREY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# Utility helpers
# ══════════════════════════════════════════════════════════════════════════════

def page_heading(doc, number, title):
    """Full-width copper page header with page number and title."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(f"Page {number}  |  ")
    _set_run(r1, size=13, bold=True, colour=COPPER)
    r2 = p.add_run(title)
    _set_run(r2, size=13, bold=True, colour=COPPER)
    # Bottom border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B87333")
    pBdr.append(bottom)
    pPr.append(pBdr)

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    body(p, text, size=12, bold=True, colour=DARK_GREY)

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    body(p, text, size=11, bold=True, italic=True, colour=DARK_GREY)

def para(doc, text, size=11, italic=False, colour=CHARCOAL, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    body(p, text, size=size, italic=italic, colour=colour)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after   = Pt(2)
    body(p, text)

def formula(doc, text):
    """Display formula in a monospace indented block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name  = "Courier New"
    r.font.size  = Pt(10)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)  # dark blue for formulas

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"▶  {text}")
    r.font.size  = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x44, 0x22)   # dark copper tone

def shade_cell(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl    = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], "2C2C2C")
        for para_ in hdr_cells[i].paragraphs:
            for run in para_.runs:
                run.font.bold  = True
                run.font.color.rgb = WHITE
                run.font.size  = Pt(9.5)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        bg    = "F5F0EB" if r_idx % 2 == 0 else "FAFAFA"
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = val
            shade_cell(cells[c_idx], bg)
            for para_ in cells[c_idx].paragraphs:
                for run in para_.runs:
                    run.font.size = Pt(9.5)

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — Big Picture
# ══════════════════════════════════════════════════════════════════════════════
page_heading(doc, 1, "Big Picture: What the Paper Is About")

para(doc,
     "The paper asks one deceptively simple question: does it matter whether you implement carry, "
     "momentum, or value as a directional (time-series) trade or as a relative-value "
     "(cross-sectional) trade? The answer is a definitive yes — and it differs by strategy.")

add_table(doc,
    headers=["Strategy", "Better Cross-Sectional?", "Better Time-Series?"],
    rows=[
        ["Value",    "Yes — mean reversion relative",       "No — overwhelmed by global trend"],
        ["Momentum", "No — hedges out strong global trend", "Yes — captures market directionality"],
        ["Carry",    "Similar performance in both",          "Similar performance in both"],
    ],
    col_widths=[4, 7, 7])

h2(doc, "The Mathematical Link")
para(doc,
     "Cross-sectional (CS) portfolio weights = time-series (TS) weights MINUS the "
     "cross-sectional average. That average is the global factor (e.g., 'all commodities rise'). "
     "So a CS portfolio is a TS portfolio with the global market trend hedged out.")

h2(doc, "Definitions")
bullet(doc, "Cross-sectional: Rank assets, go long the best, short the worst. Always market-neutral.")
bullet(doc, "Time-series: Each asset votes independently — long if signal > 0, short if signal < 0. "
            "Retains exposure to the global factor.")

h2(doc, "Practical Implication for Metals")
note(doc, "A backwardation/carry signal on copper → trade it directionally (TS: long copper futures).")
note(doc, "A value signal comparing copper to aluminium → trade it as a spread (CS: long Cu, short Al).")
note(doc, "Momentum → substantially more powerful as TS for commodities because of the China/dollar cycle.")

h2(doc, "Study Scope")
para(doc,
     "Sample: January 1990 to April 2015. Universe: 26 equity futures, 14 interest-rate swaps, "
     "31 FX pairs, 16 commodity futures. All prices in USD, mid-market. When all three strategies "
     "are combined in both TS and CS implementations, the all-asset Sharpe ratio reaches 1.61.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — Carry Foundations
# ══════════════════════════════════════════════════════════════════════════════
page_heading(doc, 2, "Economic Foundations: Carry")

h2(doc, "2.1  What Is Carry?")
para(doc,
     "Carry is the profit you earn if prices do not move — the difference between the spot "
     "price and the forward price of an asset.")
formula(doc, "Carry(t) = Spot(t) - Forward(t, T)")
para(doc,
     "A simple carry trade: buy the asset forward if carry is positive (backwardation); "
     "sell it forward if carry is negative (contango).")

h2(doc, "2.2  Why Does the Forward Price Differ from the Expected Spot?")
para(doc,
     "If forward prices were unbiased predictors of future spot prices (Eq. 1):")
formula(doc, "F(t,T) = E_t[S_T]")
para(doc,
     "By Jensen's inequality, this cannot hold simultaneously for the price and its reciprocal. "
     "Therefore forward prices are systematically biased — this bias is the engine of the carry trade.")

h2(doc, "2.3  CAPM Derivation of Carry P&L")
para(doc, "No-arbitrage forward pricing for a dividend/convenience-yield-paying asset (Eq. 4):")
formula(doc, "F(t,T) = S(t) x exp[(r - d)(T - t)]")
para(doc, "With CAPM, the expected spot price (Eq. 5):")
formula(doc, "E_t[S_T] = S(t) x exp[(r + beta*pi - d)(T - t)]")
para(doc, "where beta = asset beta, pi = equity risk premium. Expected P&L on a long forward (Eq. 7):")
formula(doc, "E_t[S_T] - F(t,T) = S(t) x exp[(r-d)(T-t)] x [exp(beta*pi*(T-t)) - 1]  > 0")
para(doc,
     "Key insight: The forward bias depends on beta*pi (market risk premium), NOT on r-d (carry). "
     "To link carry to forward bias you need a dividend/value factor in the pricing model (Eqs. 8-9), "
     "which makes the forward bias correlated with the dividend yield — the basis of the carry trade.")

h2(doc, "2.4  Commodity Carry: Backwardation")
para(doc,
     "In commodities, backwardation (spot > futures) arises from two competing theories:")
bullet(doc, "Keynes (Normal Backwardation): Producers sell futures to hedge price falls → structural "
            "excess futures supply → futures trade below expected spot to attract speculators.")
bullet(doc, "Storage Theory (Kaldor 1939, Working 1933): Low inventories → high convenience yield → "
            "high carry → backwardation. Profits persist because inventories take time to rebuild.")

note(doc, "Metals example: LME copper spot = $9,200/t, 3M future = $9,100/t. "
          "Annualised carry = 4 x (9,200/9,100 - 1) = +4.4%. Backwardated: carry signal says LONG copper.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — Momentum Foundations
# ══════════════════════════════════════════════════════════════════════════════
page_heading(doc, 3, "Economic Foundations: Momentum")

h2(doc, "3.1  What Is Momentum?")
para(doc,
     "Momentum is persistence in asset returns: past winners tend to continue outperforming, "
     "past losers continue underperforming. The empirical sweet spot is 6-12 months. "
     "Returns are mean-reverting at horizons below 1 month and above 3 years.")

h2(doc, "3.2  Risk-Based Theory")
para(doc,
     "Momentum can arise rationally: high-past-return assets may have higher expected future "
     "returns because they are more exposed to macroeconomic risk factors (e.g., industrial "
     "production growth — Liu & Zhang 2008). If realised and expected returns are highly "
     "correlated, the cross-section of winners has higher expected returns going forward.")

h2(doc, "3.3  Behavioural Model: The Disposition Effect")
para(doc,
     "Grinblatt and Han (2002) model a market with rational investors (fraction mu) and "
     "disposition investors (fraction 1-mu, who hold losers too long and sell winners early). "
     "Their demand curves (Eq. 10):")
formula(doc, "Rational:     D_t^r = 1 + beta*(F_t - P_t)")
formula(doc, "Disposition:  D_t^d = 1 + beta*(F_t - P_t) + alpha*(P_{t-1} - P_t)")
para(doc, "F_t = fundamental price, P_t = market-clearing price. alpha > 0: disposition investors "
          "demand more when today's price is below yesterday's. Aggregating demand = supply = 1 (Eq. 11):")
formula(doc, "mu*D_t^r + (1 - mu)*D_t^d = 1")
para(doc, "Solving for equilibrium price (Eq. 12):")
formula(doc, "P_t = w*F_t + (1 - w)*P_{t-1},    where w = 1 / (1 + mu*alpha)")
para(doc,
     "The equilibrium price is a weighted average of fundamental value and yesterday's price — "
     "it only partially reflects fundamentals. This underreaction creates positive autocorrelation "
     "(momentum). The smaller mu (fewer rational investors), the slower prices adjust.")
note(doc, "Example: If copper's fundamental value jumps from $9,000 to $10,000, the market price "
          "moves to only $9,200 on day 1, then $9,380 on day 2, creating a measurable uptrend "
          "that a momentum signal can exploit.")

h2(doc, "3.4  Why Momentum Works Better in Time-Series for Metals")
para(doc,
     "Commodity markets exhibit strong global co-movement driven by the dollar, risk sentiment, "
     "and China demand. When this global factor trends — e.g., a commodity supercycle — the "
     "time-series signal captures it directly. A cross-sectional signal hedges out the common "
     "trend, leaving only the idiosyncratic metal-to-metal spread, which is weaker. "
     "The paper confirms: momentum Sharpe doubles (0.42 CS to 0.96 TS) all-asset.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — Value Foundations
# ══════════════════════════════════════════════════════════════════════════════
page_heading(doc, 4, "Economic Foundations: Value")

h2(doc, "4.1  What Is Value?")
para(doc,
     "Value is the deviation of a market price from its fundamental (fair) value, exploiting "
     "mean-reversion back to fundamentals over long horizons (typically 1-5 years).")

h2(doc, "4.2  Interest Rate Value: Consumption Model")
para(doc, "From a two-period consumption model, the inter-temporal equilibrium condition (Eq. 13):")
formula(doc, "E_t [ exp(r - rho) x U'(C_{t+1}) / U'(C_t) ] = 1")
para(doc, "With log utility and log-normal consumption, the equilibrium real interest rate (Eq. 14):")
formula(doc, "r = g + rho - sigma^2")
para(doc,
     "where g = real consumption growth, rho = time preference, sigma = consumption vol. "
     "In a deterministic world: r = g. Value signal for rates (Eq. 37):")
formula(doc, "Value(t) = S_10Y(t) - GDP(t)")
para(doc,
     "A high 10Y swap rate relative to GDP growth means rates are 'expensive' (bond prices cheap) "
     "→ go long bonds (receive fixed). A low rate relative to GDP growth means go short bonds.")

h2(doc, "4.3  Equity Value: Gordon Growth / Dividend Discount")
para(doc, "Stock price as PV of dividends (Eq. 16):")
formula(doc, "P = D / (R - g_d)")
para(doc, "Implied equity yield (Eq. 17) and equity risk premium when r = g (Eq. 19):")
formula(doc, "R = g + D/P     =>    ERP = D/P")
para(doc,
     "A high dividend yield signals cheapness → go long equities. "
     "High price-to-book (Tobin's Q > 1) signals expensiveness → go short.")

h2(doc, "4.4  Commodity Value: Real Price Mean Reversion")
para(doc,
     "Commodity prices mean-revert because: (1) high prices incentivise expanded supply; "
     "(2) competitive markets pull prices toward production costs. "
     "Value in commodities = positioning for long mean-reversion cycles.")
para(doc, "Signal: CPI-deflated price relative to its own long-run expanding-window average (Eq. 36):")
formula(doc, "Value(t) = Adj_Price(t) / mean(Adj_Price),    Adj_Price(t) = Price(t) / CPI(t)")
note(doc, "Metals example: Real copper 40% above 10Y average → expensive → SHORT. "
          "Real aluminium 20% below average → cheap → LONG.")

h2(doc, "4.5  Why Value Works Better in Cross-Section")
para(doc,
     "Value relies on mean reversion to a level. Traded cross-sectionally (e.g., cheap copper "
     "vs. expensive aluminium), it isolates relative cheapness and hedges the common metal cycle. "
     "In time-series, a 'cheap copper' signal can still lose if a global commodity bear market "
     "overwhelms the mean-reversion — the global factor swamps the signal. "
     "Cross-sectional implementation eliminates this risk by construction.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — Signal Construction
# ══════════════════════════════════════════════════════════════════════════════
page_heading(doc, 5, "Signal Construction: Exact Formulas")

h2(doc, "5.1  Carry Signals")

h3(doc, "FX (Eq. 21)")
formula(doc, "Carry(t) = 4 x (Spot(t) / Fwd_3M(t) - 1)")

h3(doc, "Equity Futures — seasonality-adjusted (Eqs. 22-24)")
formula(doc, "Raw_Carry(t) = 1/(T2-T1) x [Fut(t,T1) - Fut(t,T2)] / Fut(t,T2)")
formula(doc, "Adj(t)       = mean[Raw_Carry(t) - 1Y_MA(Raw_Carry(t))]  over same calendar month")
formula(doc, "Carry(t)     = Raw_Carry(t) - Adj(t)")

h3(doc, "Commodities — 1Y curve spread (Eq. 25) [USE THIS FOR METALS]")
formula(doc, "Carry(t) = [Fut(t, 3M+12M) - Fut(t, 3M)] / Fut(t, 3M+12M)")
para(doc,
     "Using the 1-year spread (front contract vs. contract expiring 12 months later) removes "
     "seasonal effects that repeat annually. Positive = backwardation = LONG signal.",
     indent=True)

h3(doc, "Simple cash-3M proxy (when full curve unavailable)")
formula(doc, "Carry(t) = 4 x (Cash(t) / 3M(t) - 1)")

h3(doc, "Swap Rates — carry + roll-down (Eqs. 26-27)")
formula(doc, "Carry(t) = [S_10Y(t) - Fixing(t)] / Duration(t)  +  [S_10Y(t) - S_7Y(t)] / 3")
formula(doc, "Duration(t) = [1 - (1 + S_10Y(t))^(-10)] / S_10Y(t)")

h2(doc, "5.2  Momentum Signal: CTA EWMA Crossover (Eqs. 28-33)")
para(doc, "Three timescale pairs — NOT look-back days; these are EWMA speed parameters:")
formula(doc, "Short EWMAs: S_k = (8, 16, 32)    Long EWMAs: L_k = (24, 48, 96)")
para(doc, "Each number n maps to a decay factor lambda = (n-1)/n and half-life (Eq. 28):")
formula(doc, "HL = log(0.5) / log(lambda) = log(0.5) / log(1 - 1/n)")
para(doc, "Steps for each timescale k = 1, 2, 3:")
formula(doc, "(a)  x_k = EWMA[Price | S_k]  -  EWMA[Price | L_k]          (Eq. 29)")
formula(doc, "(b)  y_k = x_k / Rolling_StdDev[Price | 63 days]             (Eq. 30)")
formula(doc, "(c)  z_k = y_k / Rolling_StdDev[y_k | 252 days]             (Eq. 31)")
formula(doc, "(d)  u_k = z_k x exp(-z_k^2 / 4) / 0.89                     (Eq. 32)")
para(doc,
     "Step (d) is the non-linear response function. It peaks at |z| = sqrt(2) = 1.41 and then "
     "declines — very extreme signals are reduced, not chased. The denominator 0.89 normalises "
     "the output standard deviation to approximately 1.",
     indent=True)
formula(doc, "(e)  S_CTA = (u_1 + u_2 + u_3) / 3                          (Eq. 33)")

h2(doc, "5.3  Value Signals")
add_table(doc,
    headers=["Asset Class", "Signal Formula", "Direction"],
    rows=[
        ["FX",           "PPP deviation: log(Real_Spot) - mean(log(Real_Spot))",      "Positive = overvalued = SHORT"],
        ["Equity",       "Value(t) = Dividend_Yield(t)  (Eq. 35)",                    "High yield = cheap = LONG"],
        ["Commodity",    "Adj_Price(t) / mean(Adj_Price),  Adj_Price = Price/CPI (Eq. 36)", "Ratio > 1 = expensive = SHORT"],
        ["Swap Rates",   "Value(t) = S_10Y(t) - GDP(t)  (Eq. 37)",                   "High spread = bonds cheap = LONG"],
    ],
    col_widths=[3.5, 9, 5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 6 — Portfolio Construction
# ══════════════════════════════════════════════════════════════════════════════
page_heading(doc, 6, "Portfolio Construction")

h2(doc, "6.1  Cross-Sectional Portfolio")
bullet(doc, "At each rebalancing date, rank all assets within an asset class by signal magnitude.")
bullet(doc, "Take long positions on the 3 highest-ranked and short positions on the 3 lowest-ranked "
            "assets (6 extreme assets total).")
bullet(doc, "Use equal risk weights — position size does not depend on signal magnitude, only rank.")
bullet(doc, "Rebalance daily for most strategies.")
bullet(doc, "Exception — FX carry and FX value: trade 1/252 of the portfolio each day so that each "
            "cohort of positions is held for approximately one year, giving value time to work.")
bullet(doc, "Scale the entire portfolio to 15% annualised volatility using ex-ante vol estimates.")

h2(doc, "6.2  Time-Series Portfolio")
bullet(doc, "Within each asset class, take positions of size 1/N in all N assets simultaneously.")
bullet(doc, "Sign of position = sign of the signal (long if positive, short if negative).")
bullet(doc, "All other construction details identical to cross-sectional.")
bullet(doc, "Key difference: no ranking; you are potentially long or short every asset at once.")

h2(doc, "6.3  Signal Combination and Aggregation")
para(doc,
     "Run carry, momentum, and value separately. Combine two or three styles by computing "
     "ex-ante volatility of each and scaling to 15% before summing. When aggregating across "
     "asset classes, use equal weights (most return volatility comes from asset class, not style).")
note(doc, "Best practice: combine V + C + M. Pairwise correlations are approximately -9%, -9%, +15%, "
          "giving strong diversification benefits.")
note(doc, "The combined time-series portfolio's 3Y rolling Sharpe ratio never went negative over "
          "25 years (1990-2015).")

h2(doc, "6.4  Sizing to 15% Volatility — Practical Steps")
bullet(doc, "Compute 21-day realised volatility for each metal (in return space).")
bullet(doc, "For signal s_i in metal i with vol sigma_i, position size = s_i x (0.15 / sigma_i) x notional.")
bullet(doc, "For cross-sectional: normalise so that the long-short book's combined vol = 15% target.")
bullet(doc, "Rebalance vol estimates at least monthly; more frequently during high-vol regimes.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 7 — Empirical Results: Cross-Sectional
# ══════════════════════════════════════════════════════════════════════════════
page_heading(doc, 7, "Empirical Results: Cross-Sectional Portfolio")

h2(doc, "7.1  Sharpe Ratios by Style and Asset Class (Table 1)")
add_table(doc,
    headers=["Asset Class", "Value", "Carry", "Momentum", "Average"],
    rows=[
        ["FX",           "0.42", "0.67", "0.74",  "0.61"],
        ["Equity",       "0.39", "0.33", "0.01",  "0.24"],
        ["Commodity",    "0.07", "0.77", "0.45",  "0.43"],
        ["Interest Rate","0.56", "0.76", "-0.31", "0.34"],
        ["All Asset",    "0.75", "1.27", "0.42",  "—"],
    ],
    col_widths=[4, 3, 3, 3.5, 3.5])
para(doc,
     "Carry is the strongest standalone cross-sectional strategy (all-asset Sharpe 1.27). "
     "Momentum fails badly in equities and rates cross-sectionally. "
     "Value is consistent but weak in commodities (0.07).")

h2(doc, "7.2  Combined Signals (Table 4)")
add_table(doc,
    headers=["Combination", "All-Asset Sharpe"],
    rows=[
        ["Value + Carry",    "1.49"],
        ["Value + Momentum", "0.86"],
        ["Carry + Momentum", "1.08"],
        ["Value + Carry + Momentum", "1.40"],
    ],
    col_widths=[8, 5])
note(doc, "The V+C+M combined portfolio never lost money on a 3Y rolling basis. "
          "3Y rolling Sharpe ranged 0 to 2.30.")

h2(doc, "7.3  Return Characteristics")
add_table(doc,
    headers=["Metric", "Value", "Carry", "Momentum"],
    rows=[
        ["Monthly return skew (all-asset CS)", "+0.09",  "-0.39", "+0.10"],
        ["Max drawdown / volatility",          "-1.8x",  "-3.1x", "-2.6x"],
        ["Average market beta",                "-4%",    "+2%",   "-4%"],
    ],
    col_widths=[7, 3.5, 3.5, 3.5])
para(doc,
     "Carry has NEGATIVE skew — it is prone to occasional large losses (crash risk). "
     "Value and momentum have positive skew. All three strategies are near market-neutral.")

h2(doc, "7.4  Inter-Strategy Correlations (Table 2)")
add_table(doc,
    headers=["Pair", "Overall Correlation", "Commodity Sub-Correlation"],
    rows=[
        ["Value - Carry",    "-9.0%",  "-48% (!!!)"],
        ["Value - Momentum", "-9.2%",  "-32%"],
        ["Carry - Momentum", "+15.0%", "+51%"],
    ],
    col_widths=[5, 5.5, 6.5])
para(doc,
     "The very negative value-carry correlation in commodities (-48%) reflects a deep economic "
     "truth: high real prices (expensive = value says short) go hand in hand with high inventories "
     "which cause contango (low/negative carry). So value and carry are natural diversifiers in metals.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 8 — Time-Series vs. Cross-Section
# ══════════════════════════════════════════════════════════════════════════════
page_heading(doc, 8, "Empirical Results: Time-Series vs. Cross-Section")

h2(doc, "8.1  Time-Series Sharpe Ratios (Table 5)")
add_table(doc,
    headers=["Asset Class", "Value", "Carry", "Momentum", "Average"],
    rows=[
        ["FX",           "0.27",  "0.55", "0.72", "0.51"],
        ["Equity",       "-0.13", "0.23", "0.41", "0.17"],
        ["Commodity",    "0.22",  "0.64", "0.45", "0.44"],
        ["Interest Rate","0.48",  "0.83", "0.77", "0.69"],
        ["All Asset",    "0.28",  "1.25", "0.96", "—"],
    ],
    col_widths=[4, 3, 3, 3.5, 3.5])

h2(doc, "8.2  Head-to-Head: CS vs. TS")
add_table(doc,
    headers=["Strategy", "Cross-Section Sharpe", "Time-Series Sharpe", "Winner"],
    rows=[
        ["Value",    "0.75", "0.28", "CS wins by 2.7x"],
        ["Carry",    "1.27", "1.25", "Tie"],
        ["Momentum", "0.42", "0.96", "TS wins by 2.3x"],
    ],
    col_widths=[4, 4.5, 4.5, 4.5])

h2(doc, "8.3  Why These Differences Arise")

h3(doc, "Momentum doubles in TS")
para(doc,
     "When a global commodity bull market drives all metals up, time-series momentum goes "
     "long everything and rides the trend. Cross-sectional momentum, by construction hedging "
     "out the global factor, misses this move entirely — it only captures which metal "
     "outperforms which other metal.")

h3(doc, "Value deteriorates in TS")
para(doc,
     "Value relies on mean reversion to a level. If all asset prices trend together for years "
     "(equity bull run 1990-2015, bond bull run), a time-series value signal loses money "
     "waiting for the global factor to revert. In CS, value correctly identifies the cheapest "
     "within the group; the long-short immunises against global drift.")

h3(doc, "Signal and Asset Correlation (Table 9)")
add_table(doc,
    headers=["Asset Class", "Value Signal Corr", "Carry Signal Corr", "Momentum Signal Corr", "Asset Return Corr"],
    rows=[
        ["FX",      "48%", "23%", "40%", "24%"],
        ["Equity",  "44%", "23%", "66%", "51%"],
        ["Commodity","30%", "3%", "19%",  "7%"],
        ["IR Swap", "36%", "30%", "51%", "19%"],
    ],
    col_widths=[3.5, 4, 4, 4.5, 4.5])
para(doc,
     "Commodity asset correlations are only 7% on average — the lowest of any asset class. "
     "This means the CS portfolio hedges little genuine risk (no strong global factor to hedge), "
     "explaining why CS and TS commodity Sharpes are similar. "
     "The same logic works in reverse for equities (51% asset correlation): "
     "CS momentum hedges out a lot of genuine trend alpha.")

h2(doc, "8.4  Combined CS + TS (Table 10)")
add_table(doc,
    headers=["Portfolio", "All-Asset Sharpe"],
    rows=[
        ["CS alone (V+C+M)",  "1.40"],
        ["TS alone (V+C+M)",  "1.37"],
        ["CS + TS combined",  "1.61"],
    ],
    col_widths=[8, 5])
note(doc, "43% correlation between CS and TS — substantial diversification. "
          "3Y rolling Sharpe for combined: range 0.5 to 2.7 over 1994-2015.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 9 — Application to Metals
# ══════════════════════════════════════════════════════════════════════════════
page_heading(doc, 9, "Application Roadmap: Implementing in Metals Data")

h2(doc, "9.1  Data You Already Have")
bullet(doc, "Metals Cash and 3M.xlsx — LME cash and 3-month prices for Cu, Al, Zn, Pb, Ni, Sn.")
bullet(doc, "Metals Futures Curve.csv — multi-tenor futures prices.")
bullet(doc, "Additional needed: US CPI monthly series for value signal deflation.")

h2(doc, "9.2  Carry Signal for Metals")
para(doc, "Primary signal (Eq. 25 — 1Y spread, removes seasonal effects):")
formula(doc, "Carry(t) = [Fut(t, 3M+12M) - Fut(t, 3M)] / Fut(t, 3M+12M)")
para(doc, "Fast proxy when only cash and 3M are available:")
formula(doc, "Carry(t) = 4 x (Cash(t) / 3M(t) - 1)")
note(doc, "Positive = backwardation = LONG signal. Negative = contango = SHORT signal.")
note(doc, "Cross-sectional carry: rank all 6 LME metals by carry, long top 2, short bottom 2.")
note(doc, "Time-series carry: for each metal independently, go long if carry > 0, short if < 0.")

h2(doc, "9.3  Momentum Signal for Metals")
para(doc, "Implement the 3-timescale CTA EWMA crossover on each metal's front-future price:")
formula(doc, "For k = 1, 2, 3 (S_k = 8,16,32 ; L_k = 24,48,96):")
formula(doc, "  x_k = EWMA(price, S_k) - EWMA(price, L_k)")
formula(doc, "  y_k = x_k / rolling_std(price, 63 days)")
formula(doc, "  z_k = y_k / rolling_std(y_k, 252 days)")
formula(doc, "  u_k = z_k * exp(-z_k^2 / 4) / 0.89      [response function]")
formula(doc, "  S_momentum = (u_1 + u_2 + u_3) / 3")
note(doc, "TS momentum: go long metal if S_momentum > 0, short if < 0. "
          "CS momentum: rank all 6 metals, long top 2, short bottom 2.")
note(doc, "Given the strong China demand / dollar-cycle common factor in metals, "
          "expect TS momentum to outperform CS momentum substantially.")

h2(doc, "9.4  Value Signal for Metals")
formula(doc, "Adj_Price(t) = Price(t) / CPI(t)         [CPI-deflated price]")
formula(doc, "Value(t)     = Adj_Price(t) / expanding_mean(Adj_Price)   [Eq. 36]")
note(doc, "Ratio > 1: expensive → SHORT. Ratio < 1: cheap → LONG.")
note(doc, "Use an expanding window (not rolling) to avoid look-ahead bias. "
          "This is a slow signal — hold positions 6-12 months, trade cross-sectionally.")

h2(doc, "9.5  Expected Sharpe Ratios in Metals")
add_table(doc,
    headers=["Strategy", "Implementation", "Expected Sharpe (Commodity)"],
    rows=[
        ["Carry",    "TS or CS",  "~0.64 - 0.77"],
        ["Momentum", "TS",        "~0.45"],
        ["Value",    "CS",        "~0.07 - 0.22"],
        ["V+C+M",    "TS",        "~0.79"],
        ["V+C+M",    "CS+TS",     "~0.92 (estimated)"],
    ],
    col_widths=[4, 4.5, 8])
para(doc,
     "The low value Sharpe is expected — commodity value is the weakest standalone signal. "
     "However, it adds diversification (correlation with carry = -48% in commodities), "
     "so including it in the combined portfolio improves the Sharpe ratio and skew.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 10 — Caveats and Implementation Checklist
# ══════════════════════════════════════════════════════════════════════════════
page_heading(doc, 10, "Caveats, Risk Management, and Implementation Checklist")

h2(doc, "10.1  Critical Caveats")

h3(doc, "Selection bias")
para(doc, "Carry, momentum, and value were chosen because they have worked historically. "
          "The strategy universe is not large, but these are ex-post validated factors. "
          "The paper acknowledges this honestly — it is not claiming to find the most "
          "profitable strategy, but to characterise when each implementation is best.")

h3(doc, "No over-fitting, but some experience bias")
para(doc, "The momentum lookback (medium frequency) and value lookback (long, expanding window) "
          "were set by economic logic, not optimisation. Do not re-optimise signal parameters "
          "on the metals data — the out-of-sample degradation will be severe.")

h3(doc, "Survivorship bias")
para(doc, "Failing markets were mostly excluded from the paper's study. For LME metals this is "
          "less of an issue, but liquidity crises (tin 1985, nickel 2022) create tail risks "
          "that the historical record understates.")

h3(doc, "Transaction costs excluded")
para(doc, "The paper's Sharpe ratios assume no trading costs. For LME metals, bid-ask spreads "
          "and roll costs are real. Carry and value (slow signals, low turnover) are cheaper. "
          "Momentum is higher turnover — estimate net-of-cost Sharpe by subtracting ~0.1-0.2.")

h3(doc, "Patience required")
para(doc, "Every strategy suffers 3+ consecutive years of negative performance. "
          "Carry was negative in the mid-1990s; momentum underperformed 2011-2013; "
          "value underperformed throughout the 2010s commodity bull run. "
          "Never abandon a strategy during its drawdown if the underlying logic is intact.")

h2(doc, "10.2  Carry Crash Risk")
para(doc,
     "Carry has negative skew. For commodity carry specifically, crashes occur during broad "
     "risk-off events when all backwardated markets simultaneously flip to contango "
     "(2008 financial crisis, 2020 COVID shock). A TS carry portfolio suffers large simultaneous "
     "losses across all metals. Risk management rule: monitor cross-metal carry correlation. "
     "When it spikes above ~60%, reduce TS carry exposure or shift to CS carry implementation.")

h2(doc, "10.3  Implementation Checklist for Metals Dashboard")
bullet(doc, "DATA: Load daily cash and 3M prices for Cu, Al, Zn, Pb, Ni, Sn.")
bullet(doc, "DATA: Load full futures curve (minimum: 3M and 15M for Eq. 25 carry signal).")
bullet(doc, "DATA: Download US CPI monthly for value signal; interpolate to daily.")
bullet(doc, "CARRY: Compute curve carry = (F_15M - F_3M) / F_15M for each metal daily.")
bullet(doc, "MOMENTUM: Implement 3-timescale EWMA crossover with response function.")
bullet(doc, "VALUE: Compute CPI-deflated price relative to expanding-window mean.")
bullet(doc, "SIZING: Compute 21-day realised vol per metal; scale positions to 15% annual vol target.")
bullet(doc, "CS PORTFOLIO: Rank metals by each signal; long top 2, short bottom 2; equal risk weight.")
bullet(doc, "TS PORTFOLIO: Sign(signal) x vol-scaled position per metal.")
bullet(doc, "COMBINED: Average V, C, M signals (each scaled to 15%); report V+C+M portfolio.")
bullet(doc, "MONITORING: 3M rolling Sharpe per strategy; cross-metal signal correlations; "
            "% metals in backwardation; value z-scores (cheapness/expensiveness vs. history).")

h2(doc, "10.4  The Single Most Important Takeaway")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.left_indent  = Cm(1)
p.paragraph_format.right_indent = Cm(1)
r = p.add_run(
    "Trade carry directionally (TS) — it works equally in both implementations and is easier. "
    "Trade momentum directionally (TS) — it is twice as strong as the relative version in commodities. "
    "Trade value as a spread (CS) — it is the only strategy materially better cross-sectionally. "
    "Combine all three at a consistent 15% volatility target. "
    "Historically, this framework has delivered Sharpe ratios of ~0.8 in commodities standalone "
    "and ~1.6 across all asset classes combined."
)
r.font.bold  = True
r.font.size  = Pt(11)
r.font.color.rgb = COPPER

doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
body(p_footer,
     "Summary by Kartavya Joshi  |  "
     "Source: Baz, Granger, Harvey, Le Roux, Rattray (2015) — Man AHL",
     size=9, colour=MID_GREY, italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(DOCS_DIR, "Baz_Granger_Summary.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
