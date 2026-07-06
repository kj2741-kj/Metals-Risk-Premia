"""
generate_paper_summary_doc.py
Plain-language ~12-18 page summary of:
  Gorton, Hayashi & Rouwenhorst (2013), "The Fundamentals of Commodity Futures
  Returns", Review of Finance 17, 35-105.
Output: GHR_2013_Inventory_Paper_Summary.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_REPO_ROOT = os.path.dirname(_SCRIPT_DIR)
DOCS_DIR = os.path.join(_REPO_ROOT, "docs")

doc = Document()
s = doc.sections[0]
s.top_margin=Cm(2.0); s.bottom_margin=Cm(2.0); s.left_margin=Cm(2.3); s.right_margin=Cm(2.3)
COPPER=RGBColor(0xB8,0x73,0x33); DARK=RGBColor(0x1A,0x1A,0x1A); GREY=RGBColor(0x55,0x55,0x55); LIGHT=RGBColor(0x78,0x78,0x78)
NAVY=RGBColor(0x2A,0x3A,0x5A)

def cell_bg(c,h):
    p=c._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:color"),"auto"); sh.set(qn("w:fill"),h); p.append(sh)
def h1(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t.upper()); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=COPPER
    pr=p._p.get_or_add_pPr(); b=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
    bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"4"); bt.set(qn("w:space"),"1"); bt.set(qn("w:color"),"B87333")
    b.append(bt); pr.append(b)
def h2(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(10.5); r.font.color.rgb=DARK
def body(t,sa=6,it=False,col=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(sa); p.paragraph_format.line_spacing=1.08
    r=p.add_run(t); r.font.size=Pt(10.5); r.italic=it
    if col: r.font.color.rgb=col
def bullet(t):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Inches(0.25)
    r=p.add_run(t); r.font.size=Pt(10.5)
def kv(label,text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Inches(0.25)
    r=p.add_run(label+": "); r.bold=True; r.font.size=Pt(10.5); r2=p.add_run(text); r2.font.size=Pt(10.5)
def boxnote(title,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.left_indent=Inches(0.15)
    r=p.add_run("  "+title+"  "); r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    # shade the run via paragraph background is awkward; emulate with colored label
    r.font.color.rgb=COPPER
    p2=doc.add_paragraph(); p2.paragraph_format.space_after=Pt(6); p2.paragraph_format.left_indent=Inches(0.15)
    p2.paragraph_format.right_indent=Inches(0.15)
    r2=p2.add_run(text); r2.font.size=Pt(10); r2.italic=True; r2.font.color.rgb=NAVY
def space(p=4): doc.add_paragraph().paragraph_format.space_after=Pt(p)
def table(headers,rows,widths):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(hh,w) in enumerate(zip(headers,widths)):
        c=t.rows[0].cells[i]; c.width=Inches(w); cell_bg(c,"B87333")
        r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    for ri,row in enumerate(rows):
        fill="F5EFE8" if ri%2==0 else "FAFAFA"
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.width=Inches(widths[ci]); cell_bg(c,fill)
            r=c.paragraphs[0].add_run(v); r.font.size=Pt(9)
            if ci==0: r.bold=True
            c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER
    space(6)

# ── Cover ──
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("The Fundamentals of Commodity Futures Returns"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=COPPER
st=doc.add_paragraph(); st.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=st.add_run("A Plain-Language Summary — Inventories, the Basis, and Risk Premiums"); r.font.size=Pt(12.5); r.font.color.rgb=DARK
a=doc.add_paragraph(); a.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=a.add_run("Gary B. Gorton, Fumio Hayashi & K. Geert Rouwenhorst (2013) — Review of Finance 17, pp. 35-105"); r.font.size=Pt(10); r.font.color.rgb=GREY
a2=doc.add_paragraph(); a2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=a2.add_run("Summary prepared for the LME Copper Risk-Premia project (Prof. I. Bouchoev)"); r.font.size=Pt(9.5); r.font.color.rgb=LIGHT
space(4)
body("How to read this document: it is written for a smart reader who is NOT a commodities specialist. "
     "Every technical term is defined the first time it appears, and most ideas are followed by a small "
     "worked example using copper numbers. You can read it top to bottom in about half an hour. Sections 1 "
     "and 12 are the executive summary; the middle sections build the intuition step by step.", it=True, col=GREY)

# ── 1 ──
h1("1. The One-Minute Version")
body("This is the single most cited academic study connecting the physical commodity world (warehouses, "
     "inventories, shortages) to the financial world (the returns you earn trading commodity futures). It "
     "asks a deceptively simple question: why do some commodity futures earn positive returns over time, "
     "and why does that vary so much across commodities and across years?")
body("Its answer is one word: inventories. How much of the physical commodity is sitting in storage is the "
     "master variable. When inventories are LOW, the commodity is scarce; the futures curve tends to slope "
     "downward (cheaper in the future, called 'backwardation'); and the expected reward for holding a long "
     "futures position - the 'risk premium' - is HIGH. When inventories are HIGH, the commodity is abundant; "
     "the curve slopes upward ('contango'); and the expected reward is low or negative.")
body("The authors prove this holds in two directions at once: across commodities (a scarce commodity vs an "
     "abundant one at the same moment) and over time (the same commodity in a tight year vs a flush year). "
     "They verify it on 31 commodities over 40 years (1971-2010).")
body("Then comes the practically important twist. Inventory data are messy, slow, and incomplete - hard to "
     "trade on directly. But you do not need them, because PRICES already reveal the inventory state. The "
     "shape of the futures curve (the 'basis', what traders call carry/roll yield), the recent price trend "
     "(momentum), and volatility are all read-outs of how scarce the commodity is. Strategies built on those "
     "price signals earn large, reliable returns - precisely because they are proxies for the same "
     "inventory-driven risk premium.")
boxnote("Why this matters for our project",
     "Carry and momentum 'work' on commodities because - per this paper - they are price proxies for the "
     "inventory-driven risk premium. That is the economic justification for what our dashboard trades. And "
     "because metals are cheap and easy to store, their inventory signal is the WEAKEST of all groups - which "
     "is exactly why direct metals fundamentals are so hard, and why a price-based approach is the right one.")

# ── 2 ──
h1("2. Why Should Anyone Care? The Puzzle Being Solved")
body("For decades two facts sat uncomfortably side by side. First, a basket of commodity futures has "
     "historically earned a positive return even though a futures contract requires no up-front investment "
     "in the commodity itself - you simply post collateral and agree a price. Where does that return come "
     "from? Second, the return is wildly uneven: huge for some commodities and periods, negative for others.")
body("Earlier explanations were partial. Some appealed to a 'risk premium' paid by nervous producers to "
     "speculators; others to the cost and convenience of storage. Nobody had a single framework that "
     "produced both the shape of the futures curve AND the expected return from the same underlying cause. "
     "This paper supplies that missing link and shows the common cause is the level of inventories.")
body("It also matters because of what it implies about HOW to trade commodities. If inventories drive "
     "returns, and if prices reveal inventories, then a disciplined investor does not need a supply-and-"
     "demand forecasting team to earn the commodity risk premium - they can read it off the futures curve "
     "and the price trend. That is a powerful, and testable, claim, and it reframes carry and momentum not "
     "as mysterious 'anomalies' but as the rational compensation for bearing scarcity risk.")
h2("2.1  Two ways to slice the same idea: across commodities vs over time")
body("Throughout, keep two distinct comparisons in mind, because the paper tests both. The CROSS-SECTIONAL "
     "comparison asks: at a single moment, does a scarce commodity (say copper in a tight year) offer a "
     "higher expected return than an abundant one (say natural gas in a glut)? The TIME-SERIES comparison "
     "asks: for one commodity, does it offer a higher expected return in its own tight periods than in its "
     "own flush periods? The theory says yes to both, and the evidence broadly agrees with both. Our own "
     "copper dashboard is a time-series exercise (one metal, trading its tight vs flush states); a future "
     "multi-metal extension would add the cross-sectional dimension.")

# ── 3 ──
h1("3. The Two Classic Theories It Brings Together")
body("Commodity-futures pricing rests on two old ideas. The paper's contribution is a single model in which "
     "both operate together, so that the curve shape and the risk premium fall out as two sides of one coin.")
h2("3.1  The Theory of Storage (Kaldor 1939, Working 1949, Brennan 1958)")
body("Holding the actual physical commodity carries a hidden benefit called the 'convenience yield'. If you "
     "own copper in a warehouse, you can feed a factory, satisfy a sudden order, or avoid a costly "
     "production stoppage if supply is disrupted. A paper futures contract cannot do any of that until it "
     "delivers. The convenience yield is the dollar value of that flexibility.")
body("The central claim of this theory: the convenience yield FALLS as inventories rise, and falls at a "
     "decreasing rate. Intuitively, the first few tonnes in an empty warehouse are precious (they insure "
     "you against a shortage); once the warehouse is full, one more tonne adds almost no flexibility. Plotted "
     "against inventory, the convenience yield is a downward-sloping, convex curve - high and steep when "
     "stocks are scarce, flat and near-zero when stocks are plentiful.")
body("Storage also costs money (warehousing, insurance, financing). The net relationship between today's "
     "spot price and the futures price - the 'basis' - reflects storage cost MINUS convenience yield. When "
     "inventories are high, convenience yield is tiny, storage cost dominates, and futures trade ABOVE spot "
     "(contango). When inventories are low, convenience yield is large, and spot can trade ABOVE futures "
     "(backwardation).")
boxnote("Intuition", "Convenience yield = the comfort of having the real thing on hand. Lots in storage = "
     "little comfort value (contango). Almost none in storage = huge comfort value (backwardation).")
h2("3.2  The Theory of Normal Backwardation (Keynes 1930, Hicks 1939)")
body("This theory is about WHO trades and WHY. Producers - a copper miner, a farmer - face the risk that "
     "prices fall before they sell. To protect themselves they SELL futures today, locking in a price. That "
     "leaves the futures market lopsided: lots of natural sellers (hedgers) and not enough natural buyers.")
body("To entice speculators to step in and BUY the futures the hedgers want to sell, the futures price must "
     "be set at a discount to the price everyone expects to prevail later. That discount is the speculator's "
     "reward - the risk premium - for providing price insurance to producers. In Keynes's language the "
     "market is 'normally backwardated' because futures sit below the expected future spot price.")
boxnote("Intuition", "Hedgers buy insurance by selling futures; speculators sell that insurance by buying "
     "futures; the risk premium is the insurance fee that flows from hedgers to speculators.")
body("These two theories had never been combined into one model that derives BOTH the basis and the risk "
     "premium endogenously. That gap is what GHR fill.")

# ── 4 ──
h1("4. The Vocabulary, in Plain English")
table(["Term","What it means","Rule of thumb"],
    [["Spot price","Price to buy the physical right now","-"],
     ["Futures price","Price agreed today for delivery later","-"],
     ["Basis","Spot minus futures (the curve's slope)","Positive = backwardation; negative = contango"],
     ["Backwardation","Futures below spot (downward curve)","Signals scarcity / LOW inventory"],
     ["Contango","Futures above spot (upward curve)","Signals plenty / HIGH inventory"],
     ["Convenience yield","Hidden benefit of holding the physical","HIGH when inventory is low"],
     ["Stock-out","Inventory ~ zero; cannot carry more forward","Convenience yield & risk premium spike"],
     ["Risk premium","Expected excess return to a long position","HIGH when inventory is low"],
     ["Excess return","Futures return on posted collateral","What the strategy actually earns"],
     ["Roll yield","Gain/loss from rolling to the next contract","Trader shorthand for the basis"],
     ["Normalised inventory","Inventory / its 12-month average (I/I*)","Above 1 = flush; below 1 = tight"]],
    [1.7,4.2,3.3])

# ── 5 ──
h1("5. A Worked Example: Reading the Copper Curve")
body("Suppose LME copper spot is $9,200/tonne and the 3-month future is $9,100/tonne. The future is BELOW "
     "spot, so the curve is backwardated. Annualised, the basis is roughly 4 x (9,200/9,100 - 1) ~ +4.4%. A "
     "positive basis like this signals tight inventories: the market is paying up for metal available NOW. "
     "By the logic above, a backwardated copper curve is associated with a higher expected return to being "
     "long - and a long roll captures the curve rolling 'up to' spot as the contract nears expiry.")
body("Now flip it: spot $9,000 and the 3-month at $9,150 (contango, basis ~ -2.2% annualised). Inventories "
     "are comfortable; the market will even pay you (via storage economics) to hold metal for later. The "
     "expected return to a passive long is lower, and a long roll bleeds as the curve rolls 'down to' spot.")
boxnote("Key myth corrected","Practitioners often insist you NEED backwardation (positive roll yield) to "
     "make money long commodities. The paper shows that is wrong in theory and in data: over 1971-2010 the "
     "average basis was slightly NEGATIVE (~ -1.1%/yr, mild contango on average) yet the commodity basket "
     "still earned a positive ~5.75%/yr risk premium. Backwardation helps; it is not a precondition.")

# ── 6 ──
h1("6. The Model in Plain Words (and the Famous Curve)")
body("GHR write a compact two-period model with two characters - a representative hedger (think producer) "
     "and a speculator - who both dislike risk. It is worth understanding in words because the whole paper "
     "hangs on the picture it produces.")
kv("The hedger's problem","The hedger owns some inventory and must decide how much to carry forward to next "
   "period (which cannot be negative - you cannot borrow stock from the future) and how much to hedge by "
   "selling futures. Carrying stock costs money; hedging removes price risk but the futures price they get "
   "may embed a discount.")
kv("The speculator's problem","The speculator holds no physical; they simply choose how big a futures "
   "position to take, balancing expected reward against risk. They will only take the long side if futures "
   "are cheap enough relative to the expected future spot - i.e. if there is a risk premium.")
kv("The equilibrium","Prices adjust until the hedger's desired hedge and the speculator's desired position "
   "match. Out of this fall two numbers tied to inventory: the basis (curve shape) and the risk premium.")
h2("6.1  The crucial role of the stock-out")
body("The model's nonnegativity rule - inventory cannot go below zero - is what generates the interesting, "
     "nonlinear behaviour. When inventories are comfortably positive, holders can always shift commodity "
     "from today to tomorrow, which keeps the convenience yield near zero and the basis pinned near minus "
     "the storage cost (mild contango). But when demand is high enough that inventory is driven toward zero "
     "- a 'stock-out' - that smoothing channel breaks. The convenience yield jumps, the spot price spikes "
     "above futures (sharp backwardation), and the risk premium rises.")
body("Picture the basis on the vertical axis and inventory on the horizontal axis (the paper's Figure 2). "
     "For high inventory the line is flat and low (contango, equal to minus the storage cost). As inventory "
     "falls toward the stock-out threshold, the curve bends sharply upward. That convex, hockey-stick shape "
     "is THE prediction of the theory of storage - and the paper shows real data trace exactly that shape.")
h2("6.2  A second worked example: a copper stock-out scare")
body("Imagine a strike at a major mine plus a surge in Chinese demand. Visible LME copper stocks fall to a "
     "fraction of their normal level. Three things happen together, and they are all the same event seen "
     "from different angles. (1) The convenience yield jumps - anyone who actually holds metal can name "
     "their price to a buyer who needs it now, so spot rises far above the futures. The curve flips into "
     "steep backwardation. (2) The basis - our carry signal - turns sharply positive. (3) The expected "
     "return to being long rises, because the market is effectively paying a fat insurance premium to "
     "whoever will bear the risk of holding/financing scarce metal. A carry or momentum strategy would have "
     "been pulled LONG into exactly this high-premium state - which is why those signals earn their return. "
     "When the strike ends and stocks rebuild, the convenience yield collapses, the curve drifts back toward "
     "contango, and the premium fades.")
h2("6.3  Inventories are sticky and seasonal - and metals least of all")
body("Two further facts about inventories shape everything. First, inventories are PERSISTENT: this month's "
     "level is very close to last month's (the paper reports first-order autocorrelations above 0.85). That "
     "is why slow-moving signals - a 12-month price trend, a multi-year basis average - carry information "
     "about current scarcity. Second, many commodities have strong SEASONAL inventory patterns: natural gas "
     "stocks build in summer and draw down in winter; grain stocks peak at harvest and fall before it. The "
     "authors strip out this seasonality before testing, so they measure genuine scarcity rather than the "
     "calendar.")
body("Tellingly, METALS show the LEAST inventory seasonality of any group (their seasonal regressions have "
     "the lowest explanatory power). Metals are produced and consumed steadily through the year and are easy "
     "to store, so their inventories neither swing seasonally nor get drawn to zero often. This is yet "
     "another angle on why metals are the quietest case for inventory-driven signals.")

# ── 7 ──
h1("7. Five Things the Theory Says We Should See")
kv("Prediction 1 — basis vs inventory","Inverse and convex: as stocks fall the basis rises, gently at first "
   "then sharply near a stock-out.")
kv("Prediction 2 — risk premium vs inventory","The expected return to a long position is HIGHER when "
   "inventories are LOWER.")
kv("Prediction 3 — price signals carry inventory information","A high recent price (momentum) and a high "
   "basis both indicate low inventory, hence higher expected return. So momentum and carry should predict "
   "returns even without inventory data.")
kv("Prediction 4 — volatility","Price volatility should be higher when inventories are low (no buffer stock "
   "to absorb supply/demand shocks).")
kv("Prediction 5 — it is about inventories, not positioning","The premium is driven by the physical state "
   "of the market, not simply by which way hedgers happen to be leaning.")

# ── 8 ──
h1("8. The Data Behind the Study")
kv("Universe","31 commodities spanning metals, softs, grains, meats and energies - a deliberately broad "
   "cross-section so conclusions are not specific to one market.")
kv("Period","Inventory data from December 1969; futures-return tests run January 1971 to December 2010, "
   "about 40 years.")
kv("What is measured","The excess return to a fully-collateralised long futures position, rolled monthly "
   "(post collateral, hold the nearest contract, roll before expiry).")
kv("The inventory yard-stick (important)","Raw inventory trends upward over decades and is highly seasonal, "
   "so it cannot be compared across time directly. The authors NORMALISE it: I/I*, where I* is the trailing "
   "12-month average. Above 1 means 'more stock than normal' (flush); below 1 means 'less than normal' "
   "(tight). They also lag it one month so the signal would have been available in real time. Example: if "
   "copper stocks are 10% below their own 12-month average, I/I* = 0.90.")
kv("Honest data caveats","Inventory numbers come from many inconsistent sources, are published with a lag, "
   "get revised, and never include off-exchange or in-transit stocks. So observed inventory is a NOISY proxy "
   "for true scarcity - a recurring reason the statistical results are weaker than the theory.")

h2("8.1  How returns and the basis were actually computed (plain version)")
body("So the numbers are interpretable, here is what sits behind them. The monthly EXCESS RETURN is the "
     "percentage change in the price of the nearest futures contract that will not expire next month, "
     "measured from one month-end to the next - i.e. the gain on a collateralised long that holds the front "
     "contract and rolls before delivery. The BASIS is the percentage gap between the nearest and the "
     "next-nearest contract, annualised by the number of days between them - a clean, daily-observable proxy "
     "for the curve's slope (very close to the (F1-F2)/F1 'roll yield' our own carry sleeve uses). For the "
     "LME metals, where exchange futures histories are short, the authors impute the curve from cash and "
     "3-month forward quotes - the same cash/3-month data we hold for copper. None of these choices is "
     "exotic; the point is that both the return and the signal come straight from observable prices.")

# ── 9 ──
h1("9. Finding 1 — The Basis Really Is an Inverse, Convex Function of Inventory")
body("Splitting each commodity's history into high- and low-inventory months, the basis is reliably higher "
     "(more backwardated) when inventories are low. A flexible nonlinear regression (a cubic spline) "
     "confirms the relationship is downward-sloping and convex: it steepens sharply as inventories get "
     "tight, exactly the hockey-stick of Section 6. This is a strong confirmation of the theory of storage "
     "on the broadest commodity data set assembled to that point.")
body("Crucially for us, the steepness depends on how easy the commodity is to store:")
table(["Group","Basis-inventory slope","What it means"],
    [["Energies (hard to store)","Very steep (pooled ~ -1.86)","Small inventory swings move the curve a lot"],
     ["Metals (easy to store)","Very flat (pooled ~ -0.03)","Inventory barely moves the curve"],
     ["Copper specifically","-0.06 normally; -0.20 when stocks 25% below normal","Sensitive only when genuinely tight"]],
    [2.4,3.6,3.3])
boxnote("Why metals are the flat case","Metals are cheap and easy to warehouse, so 'normal' inventory I* is "
     "large relative to demand and true stock-outs are rare. The curve only bites when stocks fall well "
     "below normal. This is the first sign that inventory-based signals will be weakest for our asset class.")

# ── 10 ──
h1("10. Finding 2 — Low Inventory Means a Higher Risk Premium")
body("Next they regress the following month's excess return on normalised inventory. The slopes are mostly "
     "negative (low inventory -> higher return), and statistically strong once commodities are grouped. But "
     "two honest qualifications matter and are worth stating plainly:")
bullet("Returns are HARD to predict. The R-squared values are tiny and individual-commodity t-statistics are "
       "often weak. The inventory effect is real but noisy - more a tilt than a precise forecast.")
kv("The metals exception","The negative inventory->return relationship is significant and large for almost "
   "every group EXCEPT the easy-to-store Metals, where it is weak. Copper on its own shows some sensitivity, "
   "but as a group, metals are the hardest place to detect an inventory effect.")
boxnote("The most important sentence for our copper work",
     "The paper's own evidence says the inventory->return channel is WEAKEST for metals. Prof. Ilia's low "
     "expectations for metals fundamentals on public data are backed by the most authoritative study in the "
     "field - this is not pessimism, it is the consensus result.")

# ── 11 ──
h1("11. Finding 3 — Sorting Commodities by Inventory")
body("A natural trading test: each month, rank all the commodities by their normalised inventory, hold the "
     "low-inventory half and (in a long-short version) short the high-inventory half. The low-inventory "
     "basket earns more - the long-short spread is roughly +3.5% per year and statistically meaningful. As "
     "the theory predicts, the low-inventory commodities also have a high basis and strong recent returns: "
     "every signal points the same way, because they are all reading the same scarcity.")
body("The portfolio 'characteristics' make the mechanism vivid. When the authors examine WHAT lands in the "
     "low-inventory basket, those commodities have a markedly higher basis (more backwardated) and much "
     "stronger prior 12-month returns than the high-inventory basket - the differences are large and highly "
     "significant. In other words, sorting on inventory automatically tends to select high-carry, "
     "high-momentum commodities. This is the empirical bridge that makes the next finding almost inevitable: "
     "if inventory predicts returns, and inventory shows up in the basis and in past returns, then the basis "
     "and past returns should predict returns too.")

# ── 12 ──
h1("12. Finding 4 — You Can Skip the Inventory Data: Price Signals Work")
body("Because inventory is noisy and slow, the authors test signals built purely from PRICES. Each is a "
     "read-out of the inventory state, and each produces a large, statistically strong long-short return:")
table(["Sort signal","Long-short return / yr","t-stat","What it selects"],
    [["Futures basis (carry)","+10.6%","3.9","Low inventory, high momentum, high vol"],
     ["12-month futures momentum","+11.9%","4.4","Low inventory, high basis"],
     ["12-month spot momentum","+11.9%","4.2","Same family as above"],
     ["Volatility (de-meaned)","+5.4%","3.6","Low inventory, high basis, high momentum"]],
    [2.6,2.0,0.8,3.0])
body("These are not four independent edges - they are four windows onto the same thing. The high-basis and "
     "high-momentum portfolios are 0.85 correlated; both load on scarce, low-inventory commodities. In plain "
     "terms: CARRY and MOMENTUM earn their keep because they proxy the inventory-driven risk premium. That is "
     "the intellectual licence for our dashboard, which trades exactly those price signals on copper.")
body("The paper also checks the obvious worry - are these just 'hard-to-store commodities always win' in "
     "disguise? It separates the cross-sectional component (some commodities are structurally high-premium) "
     "from the time-series component (a given commodity in its own tight state). Both contribute, and the "
     "basis signal still works after removing the structural part, though more modestly. For a single-metal "
     "time-series strategy like ours, it is precisely the time-series component we are harvesting.")
boxnote("Direct line to our work","Our momentum sleeve = the paper's price-momentum signal. Our carry sleeve "
     "(the change in the F1-F2 roll yield) = a dynamic version of the paper's basis signal. We are harvesting "
     "the same economics GHR document, without needing the inventory data.")

# ── 13 ──
h1("13. Finding 5 — Who Is Trading Does NOT Predict Returns")
body("A long tradition tests the Keynes idea through 'hedging pressure' - the net short position of "
     "commercial hedgers, taken from the CFTC's Commitments-of-Traders report. The intuition: if hedgers are "
     "heavily short, the premium they pay speculators should be large, so future returns should be high.")
body("GHR find that trader positions move WITH prices and inventories at the same time (a strong "
     "contemporaneous relationship, R-squared ~11%), but do NOT predict the next period's return (predictive "
     "R-squared under 1%). Positioning describes the present; it is not a forecasting edge. Commercials add "
     "shorts as prices rise, and speculators behave like momentum traders, adding longs after run-ups.")
boxnote("Practical caution","Before over-investing in COT/positioning-based ideas, remember this result: in "
     "the most careful study, positioning data are context, not alpha.")

# ── 14 ──
h1("14. The One Prediction That Failed, and the Honest Limitations")
kv("The volatility puzzle","Theory says volatility should be higher when inventories are low. In the data it "
   "is the reverse at the portfolio level - high-inventory commodities show higher spot-price volatility. "
   "This is the paper's one acknowledged inconsistency, and the authors flag it openly.")
kv("Low predictability throughout","Every return regression has a low R-squared. The inventory effect is "
   "detectable in aggregate but would be a weak stand-alone trading rule on any single commodity.")
kv("Noisy inventories","Measurement error, publication lags, revisions and missing off-exchange stocks all "
   "blunt the inventory tests; the authors are explicit that this weakens their numerical results.")
body("Prof. Ilia's framing is exactly right: a superb read for the economic intuition, but the "
     "practical/numerical payoff from fundamentals is modest - and weakest of all for easy-to-store metals.", it=True)

# ── 15 ──
h1("15. Common Misconceptions This Paper Corrects")
kv("'You need backwardation to make money'","False. The average market was mildly contango yet still paid a "
   "positive premium. Backwardation tilts the odds; it is not required.")
kv("'Roll yield IS the return'","Roll yield (the basis) is one component and a signal of inventory; the total "
   "return also includes the change in the spot price. Confusing the two leads to bad intuition.")
kv("'Commitments-of-Traders tells you where prices go next'","Not in this study - positioning is "
   "contemporaneous, not predictive.")
kv("'Carry and momentum are unrelated alpha sources'","On commodities they are strongly related (0.85 "
   "portfolio correlation) because both proxy low inventory. Diversification across them is real but smaller "
   "than it looks at first glance.")

# ── 16 ──
h1("16. What This Means for the LME Copper Project")
kv("1. Our signals are theoretically grounded","Carry (the basis) and momentum are not data-mined curios; "
   "GHR show they are price proxies for the inventory-driven risk premium. The dashboard harvests the same "
   "economics without inventory data.")
kv("2. Metals are the hard case for fundamentals","The paper's results are weakest for easy-to-store metals "
   "on BOTH the basis-inventory and inventory-return links. This independently supports keeping fundamentals "
   "exploratory and low-expectation for copper.")
kv("3. Treat positioning as context, not signal","Hedging-pressure positions do not predict returns here.")
kv("4. Value sits outside this paper","GHR is about carry/momentum via inventories; it does not endorse a "
   "price-level mean-reversion (value) signal. Our value sleeve is a separate, weaker bet and is framed as such.")
kv("5. If we ever revisit fundamentals","GHR implies any copper inventory edge will appear only when stocks "
   "are FAR below normal - the convex, steep part of the curve. That argues for an event/regime overlay, not "
   "a smooth daily factor.")

# ── 17 FAQ ──
h1("17. Questions a Newcomer Always Asks")
h2("If low inventory means high returns, why not just always hold low-inventory commodities?")
body("Because it is a risk premium, not a free lunch. Inventory is observed with noise and a lag; returns "
     "are only weakly predictable (low R-squared); and the premium is precisely the compensation for bearing "
     "real risk - in a glut you can lose badly. It is a tilt that pays on average over time, not a "
     "certainty in any given month.")
h2("Why should a price-trend signal (momentum) tell me anything about inventories?")
body("Because inventories are persistent and change slowly. A sustained price rise typically reflects a "
     "market steadily drawing down its stocks; since stocks do not snap back overnight, the recent trend is "
     "informative about how scarce the commodity is right now. That is why a 12-month trend, not just "
     "yesterday's move, carries inventory information.")
h2("Couldn't speculators just arbitrage this premium away?")
body("No - it is compensation for taking on genuine price risk that producers want to shed. As long as "
     "hedgers need to offload risk and someone must be paid to carry it, the premium persists. It can shrink "
     "when capital floods in, and it varies a lot over time, but it is not a mispricing to be 'corrected'.")
h2("Is this the same as saying 'commodities hedge inflation'?")
body("No, that is a different question (the correlation of commodity returns with inflation). This paper is "
     "about why the risk premium varies across commodities and over time. The two ideas can both be true and "
     "do not depend on each other.")
h2("Does this guarantee carry and momentum keep working?")
body("No. They are risk premia, so they are episodic and can suffer for years when scarcity reverses or "
     "everyone crowds in. The value of the paper is that it explains WHY they have worked - which makes them "
     "more trustworthy than a pattern with no economic story, but not a guarantee.")
h2("So what is the catch for metals specifically?")
body("Easy storage. Metals rarely stock out, so their basis-inventory curve is flat and the inventory signal "
     "is weak. The edge concentrates in rare, deep-shortage episodes rather than showing up as a smooth, "
     "everyday factor - which is why we trade copper's price signals directly rather than its fundamentals.")

# ── 18 Literature ──
h1("18. Where This Paper Sits in the Literature")
body("GHR is a synthesis, and it helps to see the shoulders it stands on. From the Theory of Storage side "
     "come Kaldor (1939), Working (1949) and Brennan (1958), with later empirical work by Fama & French "
     "(1987, 1988) testing the convenience-yield idea on metals. From the Normal Backwardation side come "
     "Keynes (1930) and Hicks (1939), formalised by Hirshleifer (1988, 1990). Modern, optimisation-based "
     "storage models - Deaton & Laroque (1992) and Routledge, Seppi & Spatt (2000) - show how the "
     "convenience yield and the chance of a stock-out arise endogenously.")
body("On the empirical-returns side, GHR builds on a tradition documenting a commodity risk premium and its "
     "drivers: Bessembinder (1992), Erb & Harvey (2006), Gorton & Rouwenhorst (2006), and the momentum work "
     "of Miffre & Rallis (2007), among others. What makes GHR distinctive is twofold: it is the first to "
     "derive BOTH the basis and the risk premium from a single inventory-based model, and it assembles the "
     "broadest hand-collected inventory data set (31 commodities, 40 years) to test the whole chain of "
     "predictions at once. That combination is why it became the standard reference on inventories and "
     "commodity returns.")

# ── 19 ──
h1("19. Ten-Line Cheat Sheet")
bullet("Inventories are the master variable for commodity futures returns.")
bullet("Low inventory -> high convenience yield -> backwardation -> high risk premium. High inventory -> the reverse.")
bullet("The basis-vs-inventory curve is convex (a hockey stick); it bites hardest near a stock-out.")
bullet("Low-inventory commodities outperform high-inventory ones (inventory sort ~ +3.5%/yr).")
bullet("Price signals proxy inventory: basis/carry, momentum, and volatility all work (~+5 to +12%/yr long-short).")
bullet("Carry and momentum are ~0.85 correlated - both select scarce commodities.")
bullet("Trader positioning (COT) is contemporaneous, not predictive.")
bullet("The one failed prediction: volatility does not fall with inventory as expected.")
bullet("Everything is statistically noisy (low R-squared); inventory data are imperfect.")
bullet("Metals are the weakest case for inventory signals - so trade the price proxies, as our dashboard does.")

body("Reference: Gorton, G. B., Hayashi, F., & Rouwenhorst, K. G. (2013). The Fundamentals of Commodity "
     "Futures Returns. Review of Finance, 17(1), 35-105.", sa=2, col=LIGHT)

out=os.path.join(DOCS_DIR, "GHR_2013_Inventory_Paper_Summary.docx")
doc.save(out); print("Saved:",out)
