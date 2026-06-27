"""
generate_project_doc.py
Generates the project overview Word document for Kartavya Joshi.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins (narrower for density) ──────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.2)
section.right_margin  = Cm(2.2)

# ── Helpers ───────────────────────────────────────────────────────────────────
COPPER  = RGBColor(0xB8, 0x73, 0x33)   # #B87333 copper
DARK    = RGBColor(0x1A, 0x1A, 0x1A)   # near-black
GREY    = RGBColor(0x55, 0x55, 0x55)
LIGHT   = RGBColor(0x78, 0x78, 0x78)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="B87333", size=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"),   "single")
        bd.set(qn("w:sz"),    str(size))
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tcBorders.append(bd)
    tcPr.append(tcBorders)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = COPPER
    # bottom border via paragraph border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "B87333")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    return p

def body(text, space_after=4, italic=False, colour=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.italic    = italic
    if colour:
        run.font.color.rgb = colour
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2 + 0.2 * level)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p

def mixed_bullet(parts):
    """parts = list of (text, bold) tuples"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.bold = bold
    return p

def space(pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run("LME Copper Risk Premia Research")
run.bold = True; run.font.size = Pt(16); run.font.color.rgb = COPPER

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run("A Multi-Signal Framework for Systematic Commodity Risk Premia")
run2.italic = True; run2.font.size = Pt(10.5); run2.font.color.rgb = GREY

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(12)
run3 = p3.add_run("Kartavya Joshi  |  NYU  |  Supervised Research with Prof. Ilia Bouchoev (NYU / Hartree Partners)")
run3.font.size = Pt(9); run3.font.color.rgb = LIGHT

# Divider line
p_div = doc.add_paragraph()
pPr = p_div._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot  = OxmlElement("w:bottom")
bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8")
bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "B87333")
pBdr.append(bot); pPr.append(pBdr)
p_div.paragraph_format.space_after = Pt(8)

# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Project Overview")
body(
    "This research builds a systematic risk premia framework for LME Copper, following the "
    "decomposition methodology of Bouchoev (2024) and Baz & Granger (2015). The central thesis is "
    "that copper futures returns can be decomposed into compensated risk premia — momentum, carry, "
    "and value — each grounded in a distinct economic mechanism, and that combining uncorrelated "
    "signals from fundamentals, positioning, and sentiment can produce a diversified, robust alpha stream."
)
body(
    "The project is scoped as a pilot for LME Copper (F1–F27 term structure, 2006–2025), with a planned "
    "extension to all 10 LME base metals. A live interactive dashboard has been built in Python/Streamlit "
    "and is accessible at: metals-risk-premia-kj.streamlit.app"
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. COMPLETED: TECHNICAL SIGNAL RESEARCH
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Completed — Technical Signal Research")
body(
    "Three systematic signal families have been fully implemented, back-tested (2006–2025), and "
    "walk-forward validated. All signals are computed from the LME Copper forward curve (F1–F27); "
    "PnL is calculated from F1_continuous (ratio back-adjusted, capturing roll costs).",
    space_after=6
)

# 2A Momentum
heading2("2.1  Momentum")
mixed_bullet([("MA Crossover — best: ", True), ("MA(35,43), Lag-1.  ", False),
              ("OOS walk-forward avg Sharpe: ", True), ("+0.492 (15 windows, 2010–2025).", False)])
mixed_bullet([("Post-2022 avg OOS Sharpe: ", True), ("+0.521.  ", False),
              ("13/15 OOS windows positive; 10/15 above +0.30 threshold.", False)])
mixed_bullet([("Anchors + IS-Opt Weights: ", True),
              ("MA(10,25) + MA(35,43) + MA(63,100); max-Sharpe QP re-optimised annually. "
               "Avg OOS +0.442. Optimizer assigns ~100% weight to MA(35,43) in 9/15 windows, "
               "confirming MA(35,43) as the structural anchor.", False)])
mixed_bullet([("Baz-Granger CTA Signal: ", True),
              ("Response function u = z·exp(−z²/4)/0.8579 (Eq. 29-33). Slow-graded signal "
               "that avoids hard position flips on noise. Graded signal Sharpe comparable to MA crossover.", False)])
body("Walk-forward methodology: IS = 5yr rolling (1,260 days), OOS = 1yr (252 days). "
     "Lag-1 entry confirmed correct for momentum (signal evolves gradually; no same-day urgency).",
     italic=True, colour=LIGHT, space_after=6)

# 2B Carry
heading2("2.2  Carry")
mixed_bullet([("V4 Carry Momentum — best: ", True),
              ("20-day change in the (F1−F2)/F1 roll yield (a steepening curve signals physical "
               "tightening).  ", False),
              ("Walk-forward OOS Sharpe +0.50, net of 5bps — the strongest carry signal.  ", True),
              ("Robust to transaction costs (+0.29 at 20bps) and to execution lag; "
               "only +0.12 correlated with price-momentum, so genuinely additive.", False)])
mixed_bullet([("V3 Z-score: ", True),
              ("252-day rolling standardisation of (F1−F2)/F1. OOS Sharpe +0.40 — the clean, "
               "low-turnover runner-up. Fires only when carry is unusual vs its own history.", False)])
mixed_bullet([("V1 Roll Yield (baseline): ", True),
              ("sign of (F1−F2)/F1. OOS +0.24 — the naive level signal; superseded by V4/V3.", False)])
body("Timing correction: an earlier 'Same-Day' carry of ~0.55 was a one-day look-ahead artifact "
     "(it booked the contemporaneous return). Corrected to a true same-day execution (trade at the "
     "signal's close, earn from the next day), the naive level carry is only ~+0.10 — the curve-momentum "
     "and z-score variants are what make carry a genuine premium. Carry has no optimised level parameters; "
     "the 20-day horizon was selected and then walk-forward validated.",
     italic=True, colour=LIGHT, space_after=6)

# 2C Value
heading2("2.3  Value")
mixed_bullet([("V1 MA Reversion — copper optimal: ", True), ("F8, 5yr lookback, ±10% threshold, Lag-1.  ", False),
              ("Sharpe +0.277.  ", True),
              ("F12 (the NGL energy paper's reference contract) is weaker for copper — energy calibration sub-optimal.", False)])
mixed_bullet([("V1 is the chosen sleeve — walk-forward OOS Sharpe +0.43 ", True),
              ("(it improves out-of-sample vs its +0.28 IS, i.e. robust).", False)])
mixed_bullet([("V2 Baz-Granger Reversal: ", True),
              ("10yr lookback, IS Sharpe +0.512 — but ", False),
              ("collapses to +0.07 out-of-sample.  ", True),
              ("Its edge is concentrated in the 2020–2021 COVID window, so V2 is demoted from the "
               "default despite the higher in-sample number — a cautionary IS-vs-OOS example.", False)])
body("Same-day execution is preferred for value (the look-ahead 'same-day' variant scored −0.5 to −1.8, "
     "confirming it was an artifact). V1 F8 at a 5yr lookback is the empirically optimal copper "
     "configuration and the portfolio's value sleeve.",
     italic=True, colour=LIGHT, space_after=6)

# 2D Results Table
heading2("2.4  Summary Performance Table")
body("Walk-forward OOS Sharpe, net of 5bps transaction costs, same-day execution (no look-ahead). "
     "IS = 5yr rolling (1,260d), OOS = 1yr (252d), non-overlapping.", space_after=4)

tbl = doc.add_table(rows=7, cols=5)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Signal", "Best Variant", "Sharpe (Full)", "Pre-2022", "Post-2022"]
widths  = [Inches(1.3), Inches(2.2), Inches(1.0), Inches(0.9), Inches(0.9)]
for i, (h, w) in enumerate(zip(headers, widths)):
    cell = tbl.rows[0].cells[i]
    cell.width = w
    set_cell_bg(cell, "B87333")
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ("Momentum",    "MA(35,43)",                     "+0.49 OOS",   "~+0.60",   "+0.52"),
    ("Carry",       "Δ20d (F1−F2)/F1 momentum",      "+0.50 OOS",   "+0.49",    "+0.52"),
    ("Value V1",    "F8, 5yr, ±10%",                 "+0.43 OOS",   "~+0.30",   "+0.50"),
    ("Value V2",    "BG Reversal, 10yr (demoted)",   "+0.07 OOS",   "−0.10",    "+0.30"),
    ("EW Portfolio","1/3 Mom + 1/3 Carry + 1/3 Val", "+0.91 OOS",   "+0.98",    "+1.15"),
    ("Diversif. note","13/16 OOS windows positive",  "+1.00 full",  "—",        "—"),
]
fill_colors = ["F5EFE8", "FAFAFA", "F5EFE8", "FAFAFA", "F5EFE8", "FAFAFA"]
for ri, (row_data, fill) in enumerate(zip(rows_data, fill_colors)):
    for ci, val in enumerate(row_data):
        cell = tbl.rows[ri+1].cells[ci]
        set_cell_bg(cell, fill)
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if ci == 0:
            run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci >= 2 else WD_ALIGN_PARAGRAPH.LEFT

space(6)

# ══════════════════════════════════════════════════════════════════════════════
# 3. IN PROGRESS: FUNDAMENTAL SIGNALS
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. In Progress — Fundamental Signal Research")
body(
    "Fundamental signals aim to quantify the supply-demand imbalance in the physical copper market "
    "and its macro context. Unlike technical signals (which react to price), fundamental signals "
    "are leading or coincident indicators of the underlying commodity balance sheet. "
    "Data has been collected via Bloomberg (pdblp) across ten structured domains:",
    space_after=6
)

# Fundamental data table
fund_tbl = doc.add_table(rows=11, cols=3)
fund_tbl.style = "Table Grid"
fund_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

fheaders = ["Domain", "Key Series", "Signal Rationale"]
fwidths  = [Inches(1.5), Inches(2.2), Inches(2.6)]
for i, (h, w) in enumerate(zip(fheaders, fwidths)):
    cell = fund_tbl.rows[0].cells[i]
    cell.width = w
    set_cell_bg(cell, "B87333")
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

fund_rows = [
    ("01 Copper Physical",
     "LME & SHFE warehouse stocks, cancelled warrants, on-warrant inventory",
     "Stock-to-use ratio; cancelled warrant surge → imminent drawdown → backwardation signal"),
    ("02 Demand Growth",
     "China IP YoY, NBS/Caixin PMI, Fixed Asset Investment, US ISM, JPM Global PMI, Baltic Dry Index",
     "Forward-looking demand proxy; PMI above/below 50 as regime indicator for copper demand"),
    ("03 Currencies",
     "DXY, USD/CNY, USD/CLP (Chile), USD/PEN (Peru), USD/AUD + 9 other producer FX",
     "USD strength → copper headwind; producer FX weakness → supply disruption / lower capex"),
    ("04 Monetary & Rates",
     "US Fed Funds, 10Y Treasury, 10Y Breakeven, SOFR, China M2 YoY, SHIBOR, US/EU M2",
     "Real rate regime; China credit expansion drives ~50% of global copper demand"),
    ("05 Cross-Commodity",
     "WTI, Brent, Natural Gas, Coal, Gold, LME Zinc/Aluminium/Nickel, GSCI, BCOM",
     "Energy cost pass-through to smelting; base metals co-movement and relative value"),
    ("06 Financial Sentiment",
     "VIX (CBOE equity vol), MOVE (bond vol), COMEX Net Non-Commercial Longs",
     "Risk-on / risk-off regime filter; positioning as contrarian or trend-following overlay"),
    ("07 Trade Flows",
     "China copper imports (total, unwrought, ore/concentrate), exports, Yangshan premium",
     "Yangshan premium is a real-time indicator of Chinese import appetite vs. LME arb"),
    ("08 Geopolitical Risk",
     "Caldara-Iacoviello GPR Index (monthly), Global/US EPU Policy Uncertainty Indices",
     "Mine supply disruptions (strikes, coups in Chile/DRC/Peru) → price spikes; risk regime filter"),
    ("09 Energy Transition",
     "iShares Clean Energy ETF, Invesco Solar ETF, China NEV monthly sales, Lithium price",
     "Long-run structural demand signal; EV/solar uptick → secular copper demand uplift"),
    ("10 Country Producers",
     "Chile & Peru mine production, GDP for top-10 producers, all producer FX",
     "Supply-side: production disruptions (El Niño, strikes, permits) tighten global balance"),
]
fill_fund = ["F5EFE8", "FAFAFA"] * 5
for ri, (row_data, fill) in enumerate(zip(fund_rows, fill_fund)):
    for ci, val in enumerate(row_data):
        cell = fund_tbl.rows[ri+1].cells[ci]
        set_cell_bg(cell, fill)
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        if ci == 0:
            run.bold = True

space(4)
body(
    "Current status: data pipeline built and validated (Bloomberg pdblp). Signal construction "
    "underway — approach is to build a composite Fundamental Score combining standardised "
    "z-scores of each domain into a single long/short signal, with regime conditioning "
    "(e.g., suppress carry signal when PMI < 48 and DXY is rising).",
    italic=True, colour=LIGHT
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. PLANNED: POSITIONING SIGNALS
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Planned — Institutional Positioning Signals")
body(
    "Institutional positioning data reveals how the \"smart money\" and commodity trading advisors "
    "(CTAs) are positioned, providing a contrarian or trend-confirmation overlay on price signals. "
    "Two primary sources are planned:",
    space_after=6
)

heading2("4.1  CFTC Commitment of Traders (COT) — COMEX Copper")
mixed_bullet([("Source: ", True), ("CFTC weekly COT report (published every Friday for Tuesday positions). "
              "Already collected: COMEX_Copper_Net_NonCommercial_Longs.csv (Bloomberg: CEI1CNCN Index).", False)])
mixed_bullet([("Signal construction: ", True),
              ("Net non-commercial (speculative hedge fund) long position as a z-score of its "
               "52-week rolling distribution. Extreme long positioning → contrarian short signal; "
               "extreme short → contrarian long. Can also be used as a trend-confirmation overlay "
               "when combined with momentum.", False)])
mixed_bullet([("Additional COT decomposition: ", True),
              ("Separate managed money (hedge funds) from producer/merchant/processor hedgers "
               "(commercials). Commercial net short → producers locking in prices → backwardation. "
               "Money manager long crowding → crowded trade warning.", False)])

heading2("4.2  LME Positioning and Warrant Holdings")
mixed_bullet([("LME Banding Report: ", True),
              ("Published weekly. Shows percentage of total warrants held by the top 1-5 participants. "
               "High concentration → potential for short squeeze or delivery squeeze (historically "
               "relevant for copper: e.g., the 2021 LME nickel squeeze; copper has had similar episodes). "
               "Signal: concentration spike → backwardation premium expansion.", False)])
mixed_bullet([("SHFE Positioning: ", True),
              ("Shanghai Futures Exchange publishes weekly top-20 long/short position data by broker. "
               "Tracks Chinese speculative and hedging flows — critical given China's ~55% share of "
               "global copper consumption.", False)])

heading2("4.3  Signal Construction Approach")
bullet("Standardise each positioning series to 52-week z-scores to remove secular trend in gross positioning.")
bullet("Define \"extreme\" zones (z > +1.5 = crowded long; z < −1.5 = crowded short) as regime filters.")
bullet("Use as overlay: confirms or fades the technical + fundamental signal direction.")
bullet("Combine with momentum signal — positioning extremes tend to precede reversals, creating a natural value overlay.")

space(4)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PLANNED: SENTIMENT SIGNALS
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Planned — Sentiment Signals")
body(
    "Sentiment signals capture the behavioural and expectations-driven component of copper pricing — "
    "how fear, greed, and narrative drive short-term price dislocation from fundamental fair value. "
    "Three complementary approaches are proposed:",
    space_after=6
)

heading2("5.1  Options-Implied Sentiment (Copper Vol Surface)")
mixed_bullet([("Put-Call Skew (25-delta risk reversal): ", True),
              ("The relative cost of OTM copper put options vs OTM calls reflects hedger demand and "
               "tail-risk perception. A spike in put premium (negative skew) signals fear of sharp "
               "downside — often a contrarian buy signal at extremes.", False)])
mixed_bullet([("Implied Volatility vs Realised Volatility (VRP — Volatility Risk Premium): ", True),
              ("When IV >> RV, the market is overpaying for protection — "
               "typically a short copper signal (mean reversion after fear spike). "
               "When IV ≈ RV, the premium is fairly priced.", False)])
mixed_bullet([("Data source: ", True),
              ("CME/LME copper options. Bloomberg tickers for 1M ATM implied vol (LMCADS01V1M Index) "
               "and 25-delta risk reversal. Can also derive VRP from rolling realised vol.", False)])

heading2("5.2  Macro Fear / Risk-Off Regime (Already Collected)")
mixed_bullet([("VIX (CBOE Equity Volatility): ", True),
              ("Risk-off regime filter. VIX > 30 historically coincides with copper demand destruction "
               "and price collapse; carry and momentum signals degrade in high-VIX environments.", False)])
mixed_bullet([("MOVE Index (Bond Volatility): ", True),
              ("Interest rate uncertainty filter. High MOVE → macro uncertainty → copper risk premium "
               "compressed. Useful as a signal-dampening overlay.", False)])
mixed_bullet([("Proposed signal: ", True),
              ("Composite Fear Gauge = standardised average of VIX z-score + MOVE z-score. "
               "If gauge > 1.5 std: reduce position size by 50% across all signals (risk overlay). "
               "If gauge < −0.5 std (low volatility, risk-on): full position.", False)])

heading2("5.3  News and Analyst Sentiment (NLP / Consensus)")
mixed_bullet([("Bloomberg News Sentiment (BNS): ", True),
              ("Daily aggregated news sentiment score for copper (scale −1 to +1). "
               "Can construct a 10-day exponentially weighted moving average of sentiment "
               "as a short-term contrarian or trend signal.", False)])
mixed_bullet([("Broker Consensus Changes: ", True),
              ("Track 3-month consensus price target changes across copper mining equities "
               "(BHP, Freeport, Antofagasta, Glencore). Consensus upgrades → bullish signal; "
               "mass downgrades after a rally → crowded trade warning.", False)])
mixed_bullet([("Google Trends: ", True),
              ("Weekly search volume for 'copper price', 'LME copper', 'copper shortage' as a "
               "retail sentiment gauge. Extremes in search interest historically lag institutional moves "
               "and can identify late-cycle positioning.", False)])

space(4)

# ══════════════════════════════════════════════════════════════════════════════
# 6. PORTFOLIO CONSTRUCTION (NEXT STEP)
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Portfolio Construction — Next Step")
body(
    "Once all four signal families are built, the research will combine them into a multi-factor "
    "portfolio. The equal-weight (EW) benchmark is the natural starting point:",
    space_after=4
)

p = doc.add_paragraph()
p.paragraph_format.left_indent  = Inches(0.4)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run("Port_pos[t]  =  w₁ × Mom[t]  +  w₂ × Carry[t]  +  w₃ × Value[t]  +  w₄ × Fundamental[t]  +  w₅ × Sentiment[t]")
run.font.size = Pt(9); run.bold = True; run.font.color.rgb = COPPER

mixed_bullet([("EW baseline: ", True), ("Equal weight (1/5 each). If pairwise correlations < 0.3, "
              "EW Sharpe ≈ √5 × avg individual ≈ 1.0–1.2 (theoretical diversification bound).", False)])
mixed_bullet([("Risk parity variant: ", True), ("Inverse-volatility weighting to equalise risk contribution "
              "across signals — reduces dominance of high-vol momentum signal in volatile regimes.", False)])
mixed_bullet([("Regime-conditional weights: ", True), ("Downweight carry post-2022 (contango regime); "
              "upweight value in dislocation regimes (high VIX + low positioning z-score).", False)])
mixed_bullet([("Sub-period attribution: ", True), ("Reproduce the NGL energy paper's sub-period table across "
              "Pre-2022 / Post-2022 and extend to Pre-2020 / COVID / Post-COVID splits.", False)])
mixed_bullet([("10-metal extension: ", True), ("Apply same framework to all LME base metals "
              "(Aluminium, Zinc, Lead, Nickel, Tin, Cobalt, Molybdenum, Lithium, Steel).", False)])

space(4)

# ══════════════════════════════════════════════════════════════════════════════
# 7. TECHNOLOGY & DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Technology Stack & Live Dashboard")
mixed_bullet([("Language: ", True), ("Python (pandas, numpy, scipy, plotly, streamlit).", False)])
mixed_bullet([("Data: ", True), ("LME futures curve (F1–F27), Bloomberg Terminal (pdblp), "
              "CFTC public COT, Caldara-Iacoviello GPR Index.", False)])
mixed_bullet([("Dashboard: ", True), ("Streamlit app deployed on Streamlit Cloud — "
              "metals-risk-premia-kj.streamlit.app", False)])
mixed_bullet([("Tabs live: ", True), ("Market Overview · Term Structure · Cash vs 3M · "
              "Volume & OI · LME-CME Spread · Statistics · "
              "Momentum Signals (Tab 7) · Carry Signals (Tab 8) · Value Signals (Tab 9) · "
              "Portfolio Construction (Tab 10 — EW + inverse-vol weighting toggle)", False)])
mixed_bullet([("In progress: ", True), ("Fundamental and Sentiment sleeves (to extend the portfolio "
              "from 3 to 5 signal families)", False)])
mixed_bullet([("Repository: ", True), ("github.com/kj2741-kj/Metals-Risk-Premia (branch: main)", False)])

space(4)

# ══════════════════════════════════════════════════════════════════════════════
# 8. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading1("References")
body("Baz, J. & Granger, N. M. (2015). Dissecting Investment Strategies in the Cross Section and Time Series. SSRN 2695101.")
body("Bouchoev, I. (2024). Lecture notes on commodity risk premia. NYU.")
body("NGL Energy Risk-Premia paper. Risk Premia in Diversified Energy Portfolios.")
body("CFTC (2025). Commitments of Traders — Disaggregated Futures Only Report. cftc.gov.")
body("Caldara, D. & Iacoviello, M. (2022). Measuring Geopolitical Risk. American Economic Review 112(4).")

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"C:\Users\Kartavya\Metals Risk Premia\LME_Copper_RiskPremia_ProjectOverview.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
